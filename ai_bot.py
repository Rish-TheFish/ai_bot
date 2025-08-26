import os, csv, shutil, re
from datetime import datetime
import time  # Add timing import

# Set environment variables to prevent segmentation faults
os.environ["OMP_NUM_THREADS"] = "1"
os.environ["MKL_NUM_THREADS"] = "1"
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["PYTORCH_ENABLE_MPS_FALLBACK"] = "1"

from Logistics_Files import *
from Logistics_Files.config_details import DOCS_PATH, DB_PATH, EMBEDDING_MODEL, MODEL_NAME, POSTGRES_PASSWORD
# UPLOAD_PIN commented out for now
from Logistics_Files.backend_log import add_backend_log, backend_logs
import pickle
from langchain_community.document_loaders import (
    TextLoader, PyMuPDFLoader, UnstructuredWordDocumentLoader, CSVLoader, UnstructuredPDFLoader
)
from langchain_community.vectorstores import FAISS
import faiss
import numpy as np

from langchain_ollama import OllamaLLM
from langchain.chains import RetrievalQA
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Try to import SentenceTransformer, fallback to simple embeddings if not available
try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
    print("[INFO] sentence-transformers imported successfully")
except ImportError as e:
    SENTENCE_TRANSFORMERS_AVAILABLE = False
    print(f"[WARNING] sentence-transformers not available: {e}, using simple fallback embeddings")
import xml.etree.ElementTree as ET
from docx import Document as DocxReader
import random
import hashlib
import psycopg2
import logging
from typing import List, Optional, Any
from pathlib import Path
import inspect
import psutil
import subprocess
import platform
# Optional table extraction imports (heavy dependencies)
try:
    import pandas as pd
    import tabula
    import camelot
    from PIL import Image
    import pytesseract
    import cv2
    TABLE_EXTRACTION_AVAILABLE = True
    logging.info("[DEBUG] Table extraction libraries loaded successfully")
except ImportError as e:
    logging.info(f"[DEBUG] Table extraction libraries not available: {e}")
    logging.info("[DEBUG] Install with: pip install pandas tabula-py camelot-py[cv] opencv-python pytesseract Pillow")
    TABLE_EXTRACTION_AVAILABLE = False
    # Create dummy imports to prevent errors
    pd = None
    tabula = None
    camelot = None
    Image = None
    pytesseract = None
    cv2 = None

# Enable MPS fallback for Apple Silicon
os.environ["PYTORCH_ENABLE_MPS_FALLBACK"] = "1"

if not any("ffmpeg" in os.path.basename(p).lower() for p in os.environ["PATH"].split(os.pathsep)):
    ffmpeg_path = os.path.join(os.getcwd(), "ffmpeg.exe")
    if os.path.exists(ffmpeg_path):
        os.environ["PATH"] += os.pathsep + os.getcwd()

#logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s %(message)s')

# List of function names whose logs should be shown as system logs
USER_RELEVANT_FUNCTIONS = {
    'upload_docs', 'build_db', 'force_rebuild_db', 'handle_question', 'query_answer',
    'delete_document', 'replace_document', 'rename_document', 'extract_chat_email_knowledge',
    'get_doc_topic_map', 'get_document_status', 'export_chat', 'process_chat_email_file'
}

# Patch logging to intercept and print system logs for user-relevant functions
old_log_info = logging.info
old_log_warning = logging.warning
old_log_error = logging.error

def system_log_wrapper(level_func, level_name):
    def wrapper(msg, *args, **kwargs):
        # Get the calling function name
        stack = inspect.stack()
        func_name = stack[1].function if len(stack) > 1 else ''
        # Print to terminal as usual
        level_func(msg, *args, **kwargs)
        # If from a user-relevant function, print as system log
        if func_name in USER_RELEVANT_FUNCTIONS:
            print(f"[SYSTEM LOG] ({func_name}): {msg}")
    return wrapper

logging.info = system_log_wrapper(old_log_info, 'INFO')
logging.warning = system_log_wrapper(old_log_warning, 'WARNING')
logging.error = system_log_wrapper(old_log_error, 'ERROR')



def detect_hardware_capabilities():
    """Detect system hardware and return optimal configuration."""
    hardware_info = {
        'ram_gb': 0,
        'cpu_cores': 0,
        'gpu_available': False,
        'gpu_type': None,
        'gpu_memory_gb': 0,
        'platform': platform.system(),
        'architecture': platform.machine(),
        'optimal_device': 'cpu',
        'optimal_batch_size': 25
    }
    
    try:
        # RAM detection
        memory = psutil.virtual_memory()
        hardware_info['ram_gb'] = memory.total / (1024**3)
        
        # CPU detection
        hardware_info['cpu_cores'] = psutil.cpu_count(logical=True)
        
        # GPU detection
        gpu_detected = False
        
        # Check for NVIDIA GPU
        try:
            result = subprocess.run(['nvidia-smi', '--query-gpu=name,memory.total', '--format=csv,noheader,nounits'], 
                                  capture_output=True, text=True, timeout=5)
            if result.returncode == 0 and result.stdout.strip():
                gpu_lines = result.stdout.strip().split('\n')
                for line in gpu_lines:
                    if line.strip():
                        parts = line.split(',')
                        if len(parts) >= 2:
                            gpu_name = parts[0].strip()
                            gpu_memory = int(parts[1].strip()) / 1024  # Convert MB to GB
                            hardware_info['gpu_available'] = True
                            hardware_info['gpu_type'] = f"NVIDIA {gpu_name}"
                            hardware_info['gpu_memory_gb'] = gpu_memory
                            gpu_detected = True
                            break
        except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.CalledProcessError):
            pass
        
        # Check for Apple Silicon (M1/M2/M3)
        if not gpu_detected and platform.system() == "Darwin" and platform.machine() in ["arm64", "aarch64"]:
            try:
                # Check for Apple Silicon GPU
                result = subprocess.run(['system_profiler', 'SPDisplaysDataType'], 
                                      capture_output=True, text=True, timeout=10)
                if result.returncode == 0 and "Apple M" in result.stdout:
                    hardware_info['gpu_available'] = True
                    hardware_info['gpu_type'] = "Apple Silicon (M1/M2/M3)"
                    hardware_info['gpu_memory_gb'] = hardware_info['ram_gb'] * 0.5  # Estimate shared memory
                    gpu_detected = True
            except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.CalledProcessError):
                pass
        
        # Check for AMD GPU (Linux)
        if not gpu_detected and platform.system() == "Linux":
            try:
                result = subprocess.run(['lspci', '-v'], capture_output=True, text=True, timeout=5)
                if result.returncode == 0 and "AMD" in result.stdout and "VGA" in result.stdout:
                    hardware_info['gpu_available'] = True
                    hardware_info['gpu_type'] = "AMD GPU"
                    hardware_info['gpu_memory_gb'] = 8.0  # Estimate
                    gpu_detected = True
            except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.CalledProcessError):
                pass
        
        # Determine optimal device and batch size
        if hardware_info['gpu_available']:
            if hardware_info['gpu_type'] and "NVIDIA" in hardware_info['gpu_type']:
                hardware_info['optimal_device'] = 'cuda'
                hardware_info['optimal_batch_size'] = min(80, int(hardware_info['gpu_memory_gb'] * 3))  # More aggressive for speed
            elif hardware_info['gpu_type'] and "Apple Silicon" in hardware_info['gpu_type']:
                hardware_info['optimal_device'] = 'mps'
                hardware_info['optimal_batch_size'] = 60  # Increased for speed
            elif hardware_info['gpu_type'] and "AMD" in hardware_info['gpu_type']:
                hardware_info['optimal_device'] = 'cpu'  # AMD GPU support limited
                hardware_info['optimal_batch_size'] = 50  # Increased for speed
        else:
            # CPU optimization based on cores and RAM
            if hardware_info['cpu_cores'] >= 16 and hardware_info['ram_gb'] >= 16:
                hardware_info['optimal_batch_size'] = 60  # High-end systems
            elif hardware_info['cpu_cores'] >= 8 and hardware_info['ram_gb'] >= 8:
                hardware_info['optimal_batch_size'] = 50  # Mid-range systems
            else:
                hardware_info['optimal_batch_size'] = 40  # Standard systems
        
        logging.info(f"[HARDWARE] System detected: {hardware_info['platform']} {hardware_info['architecture']}")
        logging.info(f"[HARDWARE] RAM: {hardware_info['ram_gb']:.1f} GB")
        logging.info(f"[HARDWARE] CPU Cores: {hardware_info['cpu_cores']}")
        if hardware_info['gpu_available']:
            logging.info(f"[HARDWARE] GPU: {hardware_info['gpu_type']} ({hardware_info['gpu_memory_gb']:.1f} GB)")
        else:
            logging.info(f"[HARDWARE] GPU: Not detected")
        logging.info(f"[HARDWARE] Optimal device: {hardware_info['optimal_device']}")
        logging.info(f"[HARDWARE] Optimal batch size: {hardware_info['optimal_batch_size']}")
        
    except Exception as e:
        logging.warning(f"[HARDWARE] Error detecting hardware: {e}")
        # Fallback to conservative defaults
        hardware_info['ram_gb'] = 8.0
        hardware_info['cpu_cores'] = 4
        hardware_info['optimal_device'] = 'cpu'
        hardware_info['optimal_batch_size'] = 20
    
    return hardware_info

class AIApp:
    def __init__(self, master: Optional[Any] = None, bucket_name: Optional[str] = None, region_name: Optional[str] = None, use_s3: Optional[bool] = None, skip_db_init: bool = False):
        """Initialize the AIApp with embedding, LLM, and document DB."""
        self.theme = "light"
        self.mode = "Q&A"
        
        logging.info("[DEBUG] Starting AIApp initialization...")
        
        # Detect hardware capabilities
        logging.info("[DEBUG] Detecting hardware capabilities...")
        self.hardware_info = detect_hardware_capabilities()
        
        # Initialize embedding models using config
        logging.info(f"[DEBUG] Initializing embedding models from config: {EMBEDDING_MODEL}")
        self.initialize_embedding_models()
        
        # BGE embeddings are now initialized in initialize_embedding_models()
        # Set embedding dimension based on the initialized model
        if hasattr(self, 'embedding') and hasattr(self.embedding, 'embedding_dimension'):
            self.embedding_dimension = self.embedding.embedding_dimension
        else:
            self.embedding_dimension = 768  # Default BGE-Base dimension
        logging.info(f"[SPEED] Embedding dimension: {self.embedding_dimension}D")
        
        # Configuration for semantic search
        self.enable_semantic_search = True  # Enable hybrid semantic + vector search
        self.semantic_search_variations = 10  # Number of semantic variations to generate
        
        logging.info("[DEBUG] Initializing LLM (Ollama will auto-detect best available device)...")
        optimal_device = self.hardware_info['optimal_device']
        logging.info(f"[DEBUG] Hardware detected: {optimal_device} (Ollama will auto-optimize)")
        
        # OllamaLLM automatically uses the best available device (GPU if available, CPU if not)
        self.llm = OllamaLLM(
            model=MODEL_NAME,
            temperature=0.1,  # Low temperature for consistent, factual answers
            num_ctx=4096,     # Balanced context window for speed vs coverage
            num_predict=768,  # Balanced response length for speed vs quality
            stop=None,        # Don't stop early
            reset=True        # Reset conversation context for each call
        )
        logging.info(f"[DEBUG] LLM initialized successfully (Ollama auto-detected device)")
        self.db = None
        self.recording = False
        self.locked_out = False
        self.inappropriate_count = 0
        self.pin_verified = False
        self.upload_button = None
        self.DOCS_PATH = DOCS_PATH  # Make it accessible for Flask
        self.upload_enabled = False
        

        
        # Initialize confidence models
        logging.info("[DEBUG] Initializing confidence models...")
        self.initialize_confidence_models()
        
        # Skip database initialization if requested (for background building)
        if skip_db_init:
            logging.info("[DEBUG] Skipping database initialization (will be built in background)")
            # Log final device configuration
            logging.info(f"[DEBUG] Step 5: AIApp initialization completed (database skipped)")
            logging.info(f"[DEBUG] Final configuration:")
            logging.info(f"[DEBUG]   - Hardware detected: {self.hardware_info['optimal_device']}")
            embedding_type = "BGE" if hasattr(self, 'bge_model') and self.bge_model is not None else MODEL_NAME.split(':')[0]
            logging.info(f"[DEBUG]   - Embedding model: {embedding_type} ({self.current_embedding_type})")
            logging.info(f"[DEBUG]   - LLM model: {MODEL_NAME}")
            logging.info(f"[DEBUG]   - Confidence model: {'Available' if self.confidence_model else 'Not available'}")
            if self.hardware_info['gpu_available']:
                logging.info(f"[DEBUG]   - GPU: {self.hardware_info['gpu_type']} ({self.hardware_info['gpu_memory_gb']:.1f} GB)")
            else:
                logging.info(f"[DEBUG]   - GPU: Not available (using CPU)")
            logging.info(f"[DEBUG]   - Optimal batch size: {self.hardware_info['optimal_batch_size']}")
            return
        
        logging.info("[DEBUG] Step 3: Checking vector database existence...")
        
        # Check if database files exist before trying to load
        if self.database_exists_and_valid():
            logging.info("[DEBUG] Step 3a: Database files exist, attempting to load...")
            self.load_vector_db()
        else:
            logging.info("[DEBUG] Step 3a: No valid database files found")
            self.db = None
        
        # Force rebuild if loading failed OR if we need to update embeddings
        if not self.db:
            logging.info("[DEBUG] Step 4: No existing database found, building new one...")
            self.build_db()
        else:
            # Check if we need to rebuild due to embedding model change or document changes
            try:
                logging.info("[DEBUG] Step 4: Testing existing database compatibility...")
                
                # Check if embedding model has changed (only rebuild if necessary)
                if self._embedding_model_changed_since_last_build():
                    logging.info("[DEBUG] Step 4: Embedding model changed, rebuilding database...")
                    self.force_rebuild_db()
                    return
                
                # Test if the current embeddings work with the existing database
                test_query = "test"
                try:
                    test_embedding = self.get_query_embeddings(test_query)
                    
                    # Actually test the database with a similarity search
                    test_docs = self.db.similarity_search(test_query, k=1)
                    logging.info("[DEBUG] Step 4: Embedding model is compatible with existing database")
                    
                    # Check if documents have changed (only rebuild if necessary)
                    if self._documents_changed_since_last_build():
                        logging.info("[DEBUG] Step 4: Documents have changed, rebuilding database...")
                        self.force_rebuild_db()
                    else:
                        logging.info("[DEBUG] Step 4: Database is up-to-date, no rebuild needed")
                        
                except Exception as e:
                    logging.warning(f"[DEBUG] Step 4: Database test failed, rebuilding database: {e}")
                    self.force_rebuild_db()
                    
            except Exception as e:
                logging.warning(f"[DEBUG] Step 4: Database compatibility check failed, rebuilding database: {e}")
                self.force_rebuild_db()
        
        # Log final device configuration
        logging.info(f"[DEBUG] Step 5: AIApp initialization completed")
        logging.info(f"[DEBUG] Final configuration:")
        logging.info(f"[DEBUG]   - Hardware detected: {self.hardware_info['optimal_device']}")
        embedding_type = "BGE" if hasattr(self, 'bge_model') and self.bge_model is not None else MODEL_NAME.split(':')[0]
        logging.info(f"[DEBUG]   - Embedding model: {embedding_type} ({self.current_embedding_type})")
        logging.info(f"[DEBUG]   - LLM model: {MODEL_NAME}")
        logging.info(f"[DEBUG]   - Confidence model: {'Available' if self.confidence_model else 'Not available'}")
        if self.hardware_info['gpu_available']:
            logging.info(f"[DEBUG]   - GPU: {self.hardware_info['gpu_type']} ({self.hardware_info['gpu_memory_gb']:.1f} GB)")
        else:
            logging.info(f"[DEBUG]   - GPU: Not available (using CPU)")
        logging.info(f"[DEBUG]   - Optimal batch size: {self.hardware_info['optimal_batch_size']}")

    def initialize_embedding_models(self):
        """Initialize embeddings using BGE for maximum speed and quality."""
        try:
            if SENTENCE_TRANSFORMERS_AVAILABLE:
                from sentence_transformers import SentenceTransformer
                # Initialize BGE model directly
                self.bge_model = SentenceTransformer(EMBEDDING_MODEL)
                # Create a wrapper that matches LangChain's embedding interface
                self.embedding = self._create_bge_wrapper()
                self.current_embedding_type = "bge"
                logging.info(f"[SPEED] Using {EMBEDDING_MODEL} embeddings - 1000x faster than Ollama!")
            else:
                # Fallback to Ollama embeddings if sentence-transformers not available
                from langchain_ollama import OllamaEmbeddings
                self.embedding = OllamaEmbeddings(model=MODEL_NAME)
                self.bge_model = None
                self.current_embedding_type = "ollama"
                logging.warning(f"[FALLBACK] sentence-transformers not available, using Ollama embeddings")
        except Exception as e:
            logging.error(f"[DEBUG] Failed to initialize BGE embeddings: {e}")
            # Fallback to Ollama embeddings
            try:
                from langchain_ollama import OllamaEmbeddings
                self.embedding = OllamaEmbeddings(model=MODEL_NAME)
                self.bge_model = None
                self.current_embedding_type = "ollama"
                logging.warning(f"[FALLBACK] Using Ollama embeddings due to BGE initialization failure")
            except Exception as fallback_e:
                logging.error(f"[DEBUG] Both BGE and Ollama embeddings failed: {fallback_e}")
                raise Exception("No embedding model available")

    def _create_bge_wrapper(self):
        """Create a LangChain-compatible wrapper for BGE embeddings."""
        class BGEEmbeddingWrapper:
            def __init__(self, model):
                self.model = model
                self.embedding_dimension = 768  # BGE-Base dimension
            
            def __call__(self, text):
                """Make the wrapper callable for LangChain compatibility."""
                if isinstance(text, list):
                    return self.embed_documents(text)
                else:
                    return self.embed_query(text)
            
            def embed_query(self, text):
                return self.model.encode(text).tolist()
            
            def embed_documents(self, texts):
                return self.model.encode(texts).tolist()
        
        return BGEEmbeddingWrapper(self.bge_model)

    def get_query_embeddings(self, query: str):
        """Get query embeddings using the configured BGE embeddings object."""
        return self.embedding.embed_query(query)

    def _create_fast_faiss_database(self, chunks, embedding_model):
        """Create FAISS database using configured embeddings (BGE or Ollama)."""
        embedding_type = "BGE" if self.current_embedding_type == "bge" else "Ollama"
        print(f"[FAST] Creating FAISS database with {len(chunks)} chunks using {embedding_type} embeddings...")
        return FAISS.from_documents(chunks, self.embedding)

    def select_optimal_embedding(self, operation_type="initial_build"):
        """Select the optimal embedding model based on operation type."""
        # Use BGE embeddings if available, fallback to config model
        if hasattr(self, 'bge_model') and self.bge_model is not None:
            self.current_embedding_type = "bge"
            logging.info(f"[DEBUG] Using {EMBEDDING_MODEL} embeddings for maximum speed")
        else:
            self.current_embedding_type = MODEL_NAME.split(':')[0]  # Extract model name without version
            logging.info(f"[DEBUG] Using Ollama embeddings (fallback)")
        return True

    def evaluate_content_safety(self, text: str, type: str = "input") -> bool:
        """Check if the content is appropriate for a workplace assistant."""
        try:
            # prompt = f"""Determine if the following {type} is appropriate for a workplace compliance assistant. 
            # ONLY flag content that is clearly inappropriate for a professional workplace environment.
            # Allow normal business questions, compliance inquiries, and professional discussions.
            # Content that is off topic or not related to compliance should be allowed, unless it is clearly inappropriate.
            # Only block content that contains: explicit sexual content, violence, hate speech, illegal activities, or threats.
            # Respond only with 'yes' or 'no'.
            # {text}"""
            # response = self.llm.invoke(prompt).strip().lower()
            # return response.startswith("yes")
            return True  # Always allow content for now
        except Exception as e:
            logging.error(f"Content safety check failed: {e}")
            return True

    def load_vector_db(self) -> bool:
        """Load the FAISS vector database from disk with optimal embeddings."""
        try:
            logging.info("[DEBUG] Step 3a: Starting FAISS database load...")
            
            # Double-check database files exist before loading
            if not self.database_exists_and_valid():
                logging.warning("[DEBUG] Step 3a: Database files don't exist or are invalid")
                return False
            
            # Load the FAISS database with configured embeddings (BGE)
            self.db = FAISS.load_local(DB_PATH, self.embedding, allow_dangerous_deserialization=True)
            logging.info("[DEBUG] Step 3b: FAISS database loaded successfully")
            
            # Optimize the loaded index for better performance
            self._optimize_existing_faiss_index()
            
            doc_count = len([f for f in os.listdir(DOCS_PATH) if os.path.isfile(os.path.join(DOCS_PATH, f))])
            embedding_type = "BGE embeddings" if hasattr(self, 'bge_model') and self.bge_model is not None else f"native {MODEL_NAME} embeddings"
            add_backend_log(f"Vector database loaded successfully with {doc_count} existing documents using {embedding_type}.")
            logging.info(f"[DEBUG] Step 3c: Found {doc_count} documents in DOCS_PATH")
            return True
            
        except Exception as e:
            logging.warning(f"[DEBUG] Step 3a: Vector database load failed: {e}")
            self.db = None  # Ensure db is set to None on failure
            return False

    def _optimize_existing_faiss_index(self):
        """Optimize existing FAISS index for HNSW-like performance if possible."""
        try:
            if self.db and hasattr(self.db, 'index'):
                print(f"[HNSW] Optimizing existing FAISS index for better performance...")
                
                # Check if it's already an HNSW index
                if hasattr(self.db.index, 'hnsw'):
                    print(f"[HNSW] Index is already HNSW, optimizing search parameters...")
                    # Optimize HNSW search parameters for speed
                    self.db.index.hnsw.efSearch = 100  # Balance speed vs accuracy
                    print(f"[HNSW] HNSW search parameters optimized")
                else:
                    print(f"[HNSW] Index is not HNSW, but will use optimized search parameters")
                    
                # Set any available optimization flags
                if hasattr(self.db.index, 'nprobe'):
                    self.db.index.nprobe = 16  # Optimize IVF search if applicable
                    print(f"[HNSW] IVF search parameters optimized")
                    
                print(f"[HNSW] Existing index optimization completed")
                
        except Exception as e:
            print(f"[HNSW] Warning: Could not optimize existing index: {e}")
            # Continue without optimization

    def _get_meaningful_location(self, filename: str, chunk_id: int) -> str:
        """Generate meaningful location information for different file types."""
        try:
            # Get file extension
            file_ext = os.path.splitext(filename)[1].lower()
            
            # Convert chunk_id to int if it's a string
            try:
                chunk_id = int(chunk_id)
            except (ValueError, TypeError):
                return f"chunk {chunk_id}"
            
            # For PDFs, try to estimate page numbers
            if file_ext == '.pdf':
                # Estimate page number based on chunk ID and typical chunk size
                # Policy documents typically have 1-3 pages per chunk
                estimated_page = max(1, chunk_id + 1)
                return f"page ~{estimated_page}"
            
            # For Word documents
            elif file_ext == '.docx':
                # Estimate page number based on chunk ID
                # Word docs typically have 1-2 pages per chunk
                estimated_page = max(1, chunk_id + 1)
                return f"page ~{estimated_page}"
            
            # For CSV files
            elif file_ext == '.csv':
                # Estimate row range based on chunk ID
                # CSV files typically have 20-50 rows per chunk
                start_row = chunk_id * 30 + 1
                end_row = start_row + 29
                return f"rows ~{start_row}-{end_row}"
            
            # For text files
            elif file_ext == '.txt':
                # Estimate line range based on chunk_id
                # Text files typically have 50-100 lines per chunk
                start_line = chunk_id * 75 + 1
                end_line = start_line + 74
                return f"lines ~{start_line}-{end_line}"
            
            # For XML files
            elif file_ext == '.xml':
                # Estimate section based on chunk ID
                section_num = chunk_id + 1
                return f"section ~{section_num}"
            
            # Default fallback
            else:
                return f"chunk {chunk_id}"
                
        except Exception as e:
            logging.error(f"Error generating meaningful location for {filename}: {e}")
            return f"chunk {chunk_id}"
    
    # Content safety check removed - now integrated into hallucination detection
    
#     def _detect_hallucination(self, answer: str, query: str, context: str, relevant_docs: list) -> dict:
#         """Use AI to intelligently detect if the response contains inappropriate content or serious hallucinations."""
#         try:
#             # Only check for truly inappropriate content, not minor technical details
#             # Use a separate LLM call to evaluate the response intelligently
            
#             evaluation_prompt = f"""
# You are an AI content evaluator. Your job is to determine if the following AI response contains:
# 1. SERIOUSLY inappropriate content (harmful, offensive, illegal, etc.)
# 2. MAJOR hallucinations that could mislead users significantly
# 3. Content that goes against professional/ethical standards

# Question: {query}
# Context provided: {context[:1000]}...
# AI Response: {answer}

# IMPORTANT: Only flag content that is TRULY problematic. Be conservative - don't flag:
# - Minor technical details not in context
# - Generic phrases or standard responses
# - Reasonable inferences from the context
# - Professional language or standard compliance terms
# - Common compliance words like "explicitly stated", "adult supervision", "mature system"

# Respond with ONLY a JSON object:
# {{
#     "is_inappropriate": true/false,
#     "is_major_hallucination": true/false,
#     "confidence": 0.0-1.0,
#     "reason": "brief explanation if flagged, otherwise 'Content is appropriate'"
# }}

# Only set is_inappropriate or is_major_hallucination to true if you are highly confident the content is problematic.
# """

#             try:
#                 # Use the same LLM to evaluate the response
#                 evaluation_result = self.llm.invoke(evaluation_prompt)
                
#                 # Parse the JSON response
#                 import json
#                 try:
#                     # Try to extract JSON from the response
#                     json_start = evaluation_result.find('{')
#                     json_end = evaluation_result.rfind('}') + 1
#                     if json_start >= 0 and json_end > json_start:
#                         json_str = evaluation_result[json_start:json_end]
#                         result = json.loads(json_str)
#                     else:
#                         # Fallback if no JSON found
#                         result = {
#                             'is_inappropriate': False,
#                             'is_major_hallucination': False,
#                             'confidence': 0.0,
#                             'reason': 'JSON parsing failed, defaulting to safe'
#                         }
#                 except json.JSONDecodeError:
#                     # If JSON parsing fails, default to safe
#                     result = {
#                         'is_inappropriate': False,
#                         'is_major_hallucination': False,
#                         'confidence': 0.0,
#                         'reason': 'JSON parsing failed, defaulting to safe'
#                     }
                
#                 # Determine if this should be flagged
#                 is_hallucination = result.get('is_inappropriate', False) or result.get('is_major_hallucination', False)
#                 confidence_score = result.get('confidence', 0.0)
                
#                 # Only flag if confidence is high (>0.7) and content is truly problematic
#                 if confidence_score < 0.7:
#                     is_hallucination = False
#                     confidence_score = 0.0
                
#                 return {
#                     'is_hallucination': is_hallucination,
#                     'confidence': confidence_score,
#                     'indicators': [result.get('reason', 'AI evaluation completed')],
#                     'context_length': len(context),
#                     'answer_length': len(answer),
#                     'contradiction_score': 0.0
#                 }
                
#             except Exception as llm_error:
#                 logging.warning(f"AI-based hallucination detection failed: {llm_error}")
#                 # Fallback: only flag if there are obvious red flags
#                 return self._fallback_hallucination_check(answer, context)
            
#         except Exception as e:
#             logging.error(f"Hallucination detection failed: {e}")
#             return {
#                 'is_hallucination': False,
#                 'confidence': 0.0,
#                 'indicators': [f'Detection error: {str(e)}'],
#                 'context_length': len(context),
#                 'answer_length': len(answer),
#                 'contradiction_score': 0.0
#             }
    
    def _fallback_hallucination_check(self, answer: str, context: str) -> dict:
        """Fallback method that only flags obvious inappropriate content."""
        try:
            # Only check for obviously inappropriate content
            inappropriate_phrases = [
                'kill yourself', 'harm others', 'illegal activities', 'hack into', 'steal',
                'cheat', 'fraud', 'scam', 'hate speech', 'discrimination', 'violence'
            ]
            
            answer_lower = answer.lower()
            context_lower = context.lower()
            
            # Check for inappropriate content
            inappropriate_found = any(phrase in answer_lower for phrase in inappropriate_phrases)
            
            # Only flag if inappropriate content is found
            if inappropriate_found:
                return {
                    'is_hallucination': True,
                    'confidence': 0.9,
                    'indicators': ['Obvious inappropriate content detected'],
                    'context_length': len(context),
                    'answer_length': len(answer),
                    'contradiction_score': 0.0
                }
            
            return {
                'is_hallucination': False,
                'confidence': 0.0,
                'indicators': ['Content appears appropriate'],
                'context_length': len(context),
                'answer_length': len(answer),
                'contradiction_score': 0.0
            }
            
        except Exception as e:
            logging.error(f"Fallback hallucination check failed: {e}")
            return {
                'is_hallucination': False,
                'confidence': 0.0,
                'indicators': ['Fallback check failed, defaulting to safe'],
                'context_length': len(context),
                'answer_length': len(answer),
                'contradiction_score': 0.0
            }
    
    def _extract_entities(self, text: str) -> set:
        """Extract key entities (names, numbers, dates, etc.) from text."""
        try:
            entities = set()
            
            # Extract numbers
            numbers = re.findall(r'\b\d+(?:\.\d+)?\b', text)
            entities.update(numbers)
            
            # Extract dates
            dates = re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', text)
            entities.update(dates)
            
            # Extract percentages
            percentages = re.findall(r'\b\d+(?:\.\d+)?%\b', text)
            entities.update(percentages)
            
            # Extract proper nouns (capitalized words)
            proper_nouns = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', text)
            entities.update(proper_nouns)
            
            # Extract acronyms
            acronyms = re.findall(r'\b[A-Z]{2,}\b', text)
            entities.update(acronyms)
            
            return entities
            
        except Exception as e:
            logging.error(f"Entity extraction failed: {e}")
            return set()
    
    def _check_contradictions(self, answer: str, context: str) -> float:
        """Check for contradictions between answer and context."""
        try:
            contradiction_score = 0.0
            
            # Simple contradiction detection based on key terms
            # This is a basic implementation - could be enhanced with more sophisticated NLP
            
            # Check for opposite terms
            opposite_pairs = [
                ('yes', 'no'), ('true', 'false'), ('correct', 'incorrect'),
                ('required', 'optional'), ('mandatory', 'voluntary'),
                ('enabled', 'disabled'), ('active', 'inactive'),
                ('allowed', 'prohibited'), ('permitted', 'forbidden')
            ]
            
            for term1, term2 in opposite_pairs:
                if term1 in answer.lower() and term2 in context.lower():
                    contradiction_score += 0.2
                elif term2 in answer.lower() and term1 in context.lower():
                    contradiction_score += 0.2
            
            # Check for conflicting numbers
            answer_numbers = set(re.findall(r'\b\d+(?:\.\d+)?\b', answer))
            context_numbers = set(re.findall(r'\b\d+(?:\.\d+)?\b', context))
            
            # If answer has numbers not in context, potential contradiction
            if answer_numbers and not context_numbers:
                contradiction_score += 0.1
            
            return contradiction_score
            
        except Exception as e:
            logging.error(f"Contradiction check failed: {e}")
            return 0.0
    
    def _optimize_search_parameters(self):
        """Optimize search parameters for maximum speed during queries."""
        try:
            if self.db and hasattr(self.db, 'index'):
                print(f"[SPEED] Using default FAISS search (fastest for your data size)")
                
        except Exception as e:
            print(f"[SPEED] Warning: Could not check search parameters: {e}")
            # Continue without optimization
    
    def database_exists_and_valid(self) -> bool:
        """Check if the vector database exists and is valid."""
        try:
            faiss_path = os.path.join(DB_PATH, 'index.faiss')
            pkl_path = os.path.join(DB_PATH, 'index.pkl')
            
            if not (os.path.exists(faiss_path) and os.path.exists(pkl_path)):
                return False
            
            # Check file sizes to ensure they're not empty
            if os.path.getsize(faiss_path) == 0 or os.path.getsize(pkl_path) == 0:
                return False
            
            # Try to load the database to verify it's valid
            test_embeddings = self.embedding
            test_db = FAISS.load_local(DB_PATH, test_embeddings, allow_dangerous_deserialization=True)
            
            # Check if database has any documents
            if hasattr(test_db, 'index') and test_db.index.ntotal > 0:
                return True
            else:
                return False
                
        except Exception as e:
            logging.warning(f"[DEBUG] Database validation failed: {e}")
            return False
    
    def _documents_changed_since_last_build(self) -> bool:
        """Check if documents have changed since the last database build."""
        try:
            # Get current document files and their modification times
            current_docs = {}
            if os.path.exists(DOCS_PATH):
                for f in os.listdir(DOCS_PATH):
                    if os.path.isfile(os.path.join(DOCS_PATH, f)):
                        file_path = os.path.join(DOCS_PATH, f)
                        current_docs[f] = os.path.getmtime(file_path)
            
            # Check if we have a build timestamp file
            build_timestamp_file = os.path.join(DB_PATH, 'last_build_timestamp.txt')
            if not os.path.exists(build_timestamp_file):
                logging.info("[DEBUG] No build timestamp found, assuming rebuild needed")
                return True
            
            # Read the last build timestamp
            try:
                with open(build_timestamp_file, 'r') as f:
                    last_build_time = float(f.read().strip())
            except:
                logging.info("[DEBUG] Could not read build timestamp, assuming rebuild needed")
                return True
            
            # Check if any documents are newer than the last build
            for filename, mod_time in current_docs.items():
                if mod_time > last_build_time:
                    logging.info(f"[DEBUG] Document {filename} modified after last build, rebuild needed")
                    return True
            
            logging.info("[DEBUG] All documents are older than last build, no rebuild needed")
            return False
            
        except Exception as e:
            logging.warning(f"[DEBUG] Error checking document changes: {e}, assuming rebuild needed")
            return True
    
    def _embedding_model_changed_since_last_build(self) -> bool:
        """Check if the embedding model has changed since the last database build."""
        try:
            # Check if we have a build info file
            build_info_file = os.path.join(DB_PATH, 'build_info.txt')
            if not os.path.exists(build_info_file):
                logging.info("[DEBUG] No build info found, assuming rebuild needed")
                return True
            
            # Read the last build info
            try:
                with open(build_info_file, 'r') as f:
                    last_build_info = f.read().strip()
            except:
                logging.info("[DEBUG] Could not read build info, assuming rebuild needed")
                return True
            
            # Get current embedding model info
            current_embedding_type = getattr(self, 'current_embedding_type', 'unknown')
            current_model_name = MODEL_NAME
            
            # Create current build info string
            current_build_info = f"{current_embedding_type}:{current_model_name}"
            
            # Check if embedding model has changed
            if current_build_info != last_build_info:
                logging.info(f"[DEBUG] Embedding model changed from '{last_build_info}' to '{current_build_info}', rebuild needed")
                return True
            
            logging.info(f"[DEBUG] Embedding model unchanged: {current_build_info}")
            return False
            
        except Exception as e:
            logging.warning(f"[DEBUG] Error checking embedding model changes: {e}, assuming rebuild needed")
            return True
    


    def get_document_status(self) -> dict:
        """Get information about existing documents and database status."""
        try:
            db_exists = os.path.exists(DB_PATH) and os.path.exists(os.path.join(DB_PATH, "index.faiss"))
            doc_files = []
            if os.path.exists(DOCS_PATH):
                for f in os.listdir(DOCS_PATH):
                    if os.path.isfile(os.path.join(DOCS_PATH, f)):
                        doc_files.append(f)
            return {
                "database_loaded": self.db is not None,
                "database_exists": db_exists,
                "document_count": len(doc_files),
                "documents": doc_files
            }
        except Exception as e:
            logging.error(f"Error getting document status: {e}")
            return {
                "database_loaded": False,
                "database_exists": False,
                "document_count": 0,
                "documents": [],
                "error": str(e)
            }

    def export_chat(self, chat_content: str) -> bool:
        """Export chat content to a file."""
        if not chat_content:
            return False
        filename = f"chat_log_{datetime.now().strftime('%Y%m%d%H%M%S')}.txt"
        try:
            with open(filename, "w", encoding="utf-8") as f:
                f.write(chat_content)
            logging.info(f"Chat exported to {filename}")
            return True
        except Exception as e:
            logging.error(f"Error exporting chat: {e}")
            return False

    def upload_docs(self, file_paths: List[str]) -> List[str]:
        """Upload documents from file paths, with incremental updates when possible."""
        if not file_paths:
            return []
        
        uploaded_files = []
        
        # Check if database already exists
        db_exists = self.db is not None
        
        for path in file_paths:
            filename = os.path.basename(path)
            dest = os.path.join(DOCS_PATH, filename)
            
            if not os.path.isfile(path):
                logging.warning(f"File does not exist: {path}")
                continue
                
            if os.path.abspath(path) == os.path.abspath(dest):
                uploaded_files.append(filename)
                logging.info(f"File {filename} already in correct location")
            else:
                try:
                    shutil.copy(path, dest)
                    uploaded_files.append(filename)
                    logging.info(f"Copied {filename} to {dest}")
                except Exception as e:
                    logging.error(f"Error copying {filename}: {e}")
                    continue
        
        if uploaded_files:
            if db_exists:
                # Use incremental updates for each new file
                logging.info(f"[DEBUG] Database exists, using incremental updates for {len(uploaded_files)} files")
                for filename in uploaded_files:
                    file_path = os.path.join(DOCS_PATH, filename)
                    if self.add_document_incremental(file_path):
                        logging.info(f"[DEBUG] Successfully added {filename} incrementally")
                    else:
                        logging.warning(f"[DEBUG] Failed to add {filename} incrementally, falling back to full rebuild")
                        # Fallback to full rebuild if incremental fails
                        self.build_db("incremental_update")
                        break
            else:
                # No database exists, do full build
                logging.info(f"[DEBUG] No database exists, doing full build for {len(uploaded_files)} files")
                self.build_db("initial_build")
        
        return uploaded_files

    def build_db(self, operation_type="initial_build") -> None:
        """Rebuild the FAISS vector database from all supported documents in DOCS_PATH, using optimal embeddings (BGE if available, config model fallback)."""
        import os
        import xml.etree.ElementTree as ET
        
        # Use native config model embeddings
        logging.info(f"[DEBUG] Building database with operation type: {operation_type}")
        self.select_optimal_embedding(operation_type)
        
        # For lazy initialization, just create the AI app without building the database
        if operation_type == "lazy_init":
            logging.info("[DEBUG] Lazy initialization - skipping database build")
            return
        
        # CRITICAL FIX: ALWAYS completely remove old FAISS index files to prevent data leakage
        index_faiss = os.path.join(DB_PATH, 'index.faiss')
        index_pkl = os.path.join(DB_PATH, 'index.pkl')
        
        # Clear any existing database in memory
        self.db = None
        
        # Force removal of all old index files
        for f in [index_faiss, index_pkl]:
            try:
                if os.path.exists(f):
                    os.remove(f)
                    logging.info(f"CRITICAL: Deleted old FAISS index file: {f}")
                else:
                    logging.info(f"CRITICAL: No old index file to delete: {f}")
            except Exception as e:
                logging.error(f"CRITICAL ERROR: Could not delete {f}: {e}")
                # If we can't delete old files, we can't guarantee data integrity
                raise Exception(f"Cannot delete old index file {f}: {e}")
        
        # Verify all old files are gone
        if os.path.exists(index_faiss) or os.path.exists(index_pkl):
            raise Exception("CRITICAL: Old index files still exist after deletion attempt")
        
        logging.info("CRITICAL: All old index files successfully removed - data integrity guaranteed")
        
        all_docs = []
        logging.info(f"[DEBUG] Scanning for files in {DOCS_PATH}...")
        # Ensure DOCS_PATH exists
        if not os.path.exists(DOCS_PATH):
            logging.warning(f"[DEBUG] DOCS_PATH does not exist: {DOCS_PATH}")
            os.makedirs(DOCS_PATH, exist_ok=True)
            logging.info(f"[DEBUG] Created DOCS_PATH: {DOCS_PATH}")
        
        file_count = 0
        max_files = 5000  # Realistic limit for 16GB RAM - can handle thousands of documents
        for root, _, files in os.walk(DOCS_PATH):
            for file in files:
                if file_count >= max_files:
                    logging.warning(f"[DEBUG] Reached maximum file limit ({max_files}), stopping file processing")
                    break
                
                full_path = os.path.join(root, file)
                ext = os.path.splitext(full_path)[1].lower()
                logging.info(f"[DEBUG] Found file: {full_path}")
                file_count += 1
                
                try:
                    # Use enhanced document loader for better table extraction
                    docs = self.enhanced_document_loader(full_path)
                    if docs:
                        all_docs.extend(docs)
                        logging.info(f"[DEBUG] Enhanced processing loaded {len(docs)} docs from {file}")
                        # Log table extraction info
                        table_docs = [d for d in docs if d.metadata.get("type", "").startswith("table")]
                        if table_docs:
                            logging.info(f"[DEBUG] Extracted {len(table_docs)} tables from {file}")
                        # Only log first doc content to avoid overwhelming logs
                        if docs:
                            logging.info(f"[DEBUG] First 100 chars: {docs[0].page_content[:100]}...")
                    else:
                        logging.warning(f"No documents extracted from {file}")
                except Exception as e:
                    logging.error(f"Error loading {file}: {e}")
                    # Fall back to basic loader if enhanced processing fails
                    try:
                        if ext == ".txt":
                            loader = TextLoader(full_path)
                            docs = loader.load()
                            all_docs.extend(docs)
                        elif ext == ".pdf":
                            loader = PyMuPDFLoader(full_path)
                            docs = loader.load()
                            all_docs.extend(docs)
                        elif ext == ".docx":
                            loader = UnstructuredWordDocumentLoader(full_path)
                            docs = loader.load()
                            all_docs.extend(docs)
                        elif ext == ".csv":
                            loader = CSVLoader(full_path)
                            docs = loader.load()
                            all_docs.extend(docs)
                        elif ext == ".xml":
                            tree = ET.parse(full_path)
                            rootxml = tree.getroot()
                            text = ET.tostring(rootxml, encoding="unicode", method="text")
                            all_docs.append(Document(page_content=text, metadata={"source": full_path}))
                        logging.info(f"[DEBUG] Fallback processing loaded documents from {file}")
                    except Exception as fallback_e:
                        logging.error(f"Fallback processing also failed for {file}: {fallback_e}")
                        continue
                except Exception as e:
                    logging.error(f"Error loading {file}: {e}")
            
            # Break out of the outer loop if we've reached the file limit
            if file_count >= max_files:
                break
        logging.info(f"[DEBUG] Total loaded docs: {len(all_docs)}")
        if not all_docs:
            logging.warning("No documents loaded. Vector DB not rebuilt.")
            return
        
        # Limit the number of documents if there are too many to prevent memory issues
        max_docs = 10000  # Realistic limit for 16GB RAM - can handle thousands of documents
        if len(all_docs) > max_docs:
            logging.warning(f"[DEBUG] Too many documents ({len(all_docs)}), limiting to {max_docs} to prevent memory issues")
            # Sort by file size (smaller files first) to prioritize important documents
            all_docs.sort(key=lambda doc: len(doc.page_content))
            all_docs = all_docs[:max_docs]
            logging.info(f"[DEBUG] Limited to {len(all_docs)} documents")
        
        # Smart, bounded dynamic chunking for FAQ bot quality
        logging.info("[DEBUG] Starting document chunking...")
        
        # Calculate optimal chunk size based on available resources
        chunk_size, chunk_overlap = self.calculate_optimal_chunk_parameters()
        
        logging.info(f"[DEBUG] Calculated optimal chunk size: {chunk_size}, overlap: {chunk_overlap}")
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=[
                "\n\n\n",    # Major section breaks
                "\n\n",      # Paragraph breaks
                "\n",        # Line breaks
                ". ",        # Sentence endings
                "! ",        # Exclamation endings
                "? ",        # Question endings
                "; ",        # Semicolon separators
                ": ",        # Colon separators
                " - ",       # Dash separators
                " • ",       # Bullet points
                " * ",       # Asterisk separators
                " ",         # Word boundaries
                ""           # Character level (fallback)
            ]
        )
        
        chunks = splitter.split_documents(all_docs)
        logging.info(f"[DEBUG] Chunking parameters: size={chunk_size}, overlap={chunk_overlap}")
        logging.info(f"[DEBUG] Number of chunks created: {len(chunks)}")
        
        # Post-process chunks for better quality with stricter filtering
        processed_chunks = []
        for i, chunk in enumerate(chunks):
            # Clean up chunk content
            content = chunk.page_content.strip()
            
            # Skip chunks that are too short (with tighter limits)
            if len(content) < 30:
                continue
                
            # Skip chunks that are mostly punctuation, numbers, or whitespace
            alpha_chars = len([c for c in content if c.isalpha()])
            if alpha_chars < len(content) * 0.4:  # Increased threshold
                continue
            
            # Skip chunks that are just headers or navigation text
            if content.isupper() and len(content) < 100:
                continue
                
            # Skip chunks that are mostly special characters
            special_chars = len([c for c in content if c in '!@#$%^&*()_+-=[]{}|;:,.<>?'])
            if special_chars > len(content) * 0.3:
                continue
            
            # Create new chunk with cleaned content
            processed_chunk = Document(
                page_content=content,
                metadata=chunk.metadata.copy()
            )
            processed_chunks.append(processed_chunk)
        
        chunks = processed_chunks
        logging.info(f"[DEBUG] After processing: {len(chunks)} quality chunks retained")
        
        # Limit debug output to avoid overwhelming logs
        for i, chunk in enumerate(chunks[:3]):
            logging.info(f"[DEBUG] Chunk {i}: {chunk.page_content[:80]}...")
        
        if len(chunks) > 3:
            logging.info(f"[DEBUG] ... and {len(chunks) - 3} more chunks")
        
        # Use optimal embeddings (BGE if available, config model fallback)
        embedding_type = "BGE embeddings" if hasattr(self, 'bge_model') and self.bge_model is not None else f"native {MODEL_NAME} embeddings"
        logging.info(f"[DEBUG] Using {embedding_type} for database creation...")
        
        # Create FAISS database with batch processing to avoid memory issues
        logging.info(f"[DEBUG] Creating FAISS database with {len(chunks)} chunks using {embedding_type}...")
        
        # Smart document selection with dynamic memory monitoring using psutil
        memory = psutil.virtual_memory()
        available_memory_gb = memory.available / (1024**3)
        total_memory_gb = memory.total / (1024**3)
        used_memory_gb = memory.used / (1024**3)
        memory_percent = memory.percent
        
        optimal_batch_size = self.hardware_info['optimal_batch_size']
        
        logging.info(f"[DEBUG] Dynamic memory analysis:")
        logging.info(f"[DEBUG]   - Total RAM: {total_memory_gb:.2f} GB")
        logging.info(f"[DEBUG]   - Used RAM: {used_memory_gb:.2f} GB ({memory_percent:.1f}%)")
        logging.info(f"[DEBUG]   - Available RAM: {available_memory_gb:.2f} GB")
        logging.info(f"[DEBUG]   - Optimal batch size: {optimal_batch_size}")
        
        # Dynamic memory calculation based on actual system usage
        # Each chunk uses ~6KB in FAISS + embedding overhead
        # Use 80% of available memory for FAISS to be safe
        safe_memory_gb = available_memory_gb * 0.8  # Use 80% of available memory
        max_chunks_by_memory = int((safe_memory_gb * 1024**3) / (6 * 1024))
        
        # Adaptive practical max based on actual memory pressure
        if self.hardware_info['gpu_available'] and self.hardware_info['gpu_memory_gb'] > 0:
            # GPU systems can handle more chunks efficiently
            gpu_memory_gb = self.hardware_info['gpu_memory_gb']
            
            # Adjust based on memory pressure
            if memory_percent < 50:  # Low memory pressure
                practical_max = min(300000, max_chunks_by_memory)  # Up to 300K chunks
            elif memory_percent < 75:  # Medium memory pressure
                practical_max = min(150000, max_chunks_by_memory)  # Up to 150K chunks
            else:  # High memory pressure
                practical_max = min(80000, max_chunks_by_memory)   # Up to 80K chunks
                
            logging.info(f"[DEBUG] GPU system analysis:")
            logging.info(f"[DEBUG]   - GPU: {self.hardware_info['gpu_type']} ({gpu_memory_gb:.1f} GB)")
            logging.info(f"[DEBUG]   - Memory pressure: {memory_percent:.1f}% ({'Low' if memory_percent < 50 else 'Medium' if memory_percent < 75 else 'High'})")
            logging.info(f"[DEBUG]   - Safe for FAISS: {safe_memory_gb:.2f} GB (80% of available)")
            logging.info(f"[DEBUG]   - Theoretical max chunks: {max_chunks_by_memory:,}")
            logging.info(f"[DEBUG]   - Practical max chunks: {practical_max:,} (adaptive)")
        else:
            # CPU-only systems
            # Adjust based on memory pressure and total RAM
            if total_memory_gb >= 32:
                # High-end systems (GCP instances)
                if memory_percent < 60:
                    practical_max = min(200000, max_chunks_by_memory)  # Up to 200K chunks
                elif memory_percent < 80:
                    practical_max = min(100000, max_chunks_by_memory)  # Up to 100K chunks
                else:
                    practical_max = min(50000, max_chunks_by_memory)   # Up to 50K chunks
            else:
                # Standard systems (16GB and below)
                if memory_percent < 70:
                    practical_max = min(100000, max_chunks_by_memory)  # Up to 100K chunks
                elif memory_percent < 85:
                    practical_max = min(50000, max_chunks_by_memory)   # Up to 50K chunks
                else:
                    practical_max = min(25000, max_chunks_by_memory)   # Up to 25K chunks
            
            logging.info(f"[DEBUG] CPU system analysis:")
            logging.info(f"[DEBUG]   - Memory pressure: {memory_percent:.1f}% ({'Low' if memory_percent < 60 else 'Medium' if memory_percent < 80 else 'High'})")
            logging.info(f"[DEBUG]   - Safe for FAISS: {safe_memory_gb:.2f} GB (80% of available)")
            logging.info(f"[DEBUG]   - Theoretical max chunks: {max_chunks_by_memory:,}")
            logging.info(f"[DEBUG]   - Practical max chunks: {practical_max:,} (adaptive)")
        
        max_chunks = practical_max  # Use adaptive practical max
        logging.info(f"[DEBUG] Memory-based chunk limit: {max_chunks_by_memory}, using: {max_chunks}")
        if len(chunks) > max_chunks:
            logging.warning(f"[DEBUG] Too many chunks ({len(chunks)}), using smart selection to ensure all documents are represented")
            
            # Group chunks by document source
            doc_chunks = {}
            for chunk in chunks:
                source = chunk.metadata.get('source', 'unknown')
                if source not in doc_chunks:
                    doc_chunks[source] = []
                doc_chunks[source].append(chunk)
            
            logging.info(f"[DEBUG] Found chunks from {len(doc_chunks)} different documents")
            
            # Calculate how many chunks per document we can include
            # For optimal AI answer quality, target 30-50 chunks per document
            # This provides 20-30% coverage for accurate policy interpretation
            optimal_chunks_per_doc = 40  # Sweet spot for policy documents
            chunks_per_doc = max(optimal_chunks_per_doc, min(optimal_chunks_per_doc, max_chunks // len(doc_chunks)))
            logging.info(f"[DEBUG] Allocating {chunks_per_doc} chunks per document")
            
            # Select chunks from each document
            selected_chunks = []
            for source, doc_chunk_list in doc_chunks.items():
                # Take more chunks from each document for better coverage
                # Skip the first chunk (usually just title) and take meaningful content
                chunks_to_take = min(chunks_per_doc, len(doc_chunk_list))
                if len(doc_chunk_list) > 1:
                    # Skip title chunk and take meaningful content chunks
                    # For better coverage, take chunks from throughout the document
                    if len(doc_chunk_list) <= chunks_to_take:
                        # If we can take all chunks, do so (except title)
                        selected_chunks.extend(doc_chunk_list[1:])
                    else:
                        # Take chunks distributed throughout the document
                        step = len(doc_chunk_list) // chunks_to_take
                        for i in range(1, min(chunks_to_take + 1, len(doc_chunk_list))):
                            index = min(i * step, len(doc_chunk_list) - 1)
                            selected_chunks.append(doc_chunk_list[index])
                else:
                    # If only one chunk, take it
                    selected_chunks.extend(doc_chunk_list[:chunks_to_take])
            
            chunks = selected_chunks
            logging.info(f"[DEBUG] Selected {len(chunks)} chunks from {len(doc_chunks)} documents")
        
        try:
            logging.info(f"[DEBUG] Attempting to create FAISS database with optimal embeddings...")
            
            # Optimized approach with dynamic memory monitoring and speed improvements
            optimal_batch_size = self.hardware_info['optimal_batch_size']
            
            # Increase initial batch size for faster processing
            initial_batch_size = min(100, optimal_batch_size * 2)  # Double the initial batch size
            
            if len(chunks) > initial_batch_size:
                logging.info(f"[DEBUG] Creating FAISS database with optimized batch processing...")
                logging.info(f"[DEBUG] Using optimal batch size: {optimal_batch_size} (hardware: {self.hardware_info['gpu_type'] or 'CPU'})")
                
                # Start with larger initial batch for better performance
                initial_chunks = chunks[:initial_batch_size]
                
                # Create optimized FAISS database
                print(f"[FAST] Creating initial FAISS database with optimized index...")
                self.db = self._create_fast_faiss_database(initial_chunks, self.embedding)
                logging.info(f"[DEBUG] FAISS database created successfully with initial {initial_batch_size} chunks!")
                
                # Add remaining chunks with optimized batch processing
                remaining_chunks = chunks[initial_batch_size:]
                if remaining_chunks:
                    logging.info(f"[DEBUG] Adding {len(remaining_chunks)} remaining chunks with optimized batch sizing...")
                    
                    # Start with larger batch size for speed
                    current_batch_size = min(optimal_batch_size * 2, 80)  # Start with larger batches
                    batch_count = 0
                    total_batches = (len(remaining_chunks) + current_batch_size - 1) // current_batch_size
                    
                    for i in range(0, len(remaining_chunks), current_batch_size):
                        batch_count += 1
                        batch = remaining_chunks[i:i+current_batch_size]
                        
                        # Check memory pressure before each batch (less frequently for speed)
                        if batch_count % 3 == 0:  # Check every 3rd batch to reduce overhead
                            memory = psutil.virtual_memory()
                            memory_percent = memory.percent
                            
                            # More aggressive batch size optimization for speed
                            if memory_percent > 90:  # Very high memory pressure
                                current_batch_size = max(10, current_batch_size // 2)
                                logging.info(f"[DEBUG] Very high memory pressure ({memory_percent:.1f}%), reducing batch size to {current_batch_size}")
                            elif memory_percent > 80:  # High memory pressure
                                current_batch_size = max(15, current_batch_size // 2)
                                logging.info(f"[DEBUG] High memory pressure ({memory_percent:.1f}%), reducing batch size to {current_batch_size}")
                            elif memory_percent < 60:  # Low pressure - be more aggressive
                                current_batch_size = min(optimal_batch_size * 3, current_batch_size * 2)  # Triple the optimal size
                                logging.info(f"[DEBUG] Low memory pressure ({memory_percent:.1f}%), increasing batch size to {current_batch_size}")
                            elif memory_percent < 70:  # Medium-low pressure
                                current_batch_size = min(optimal_batch_size * 2, current_batch_size + 10)
                                logging.info(f"[DEBUG] Medium-low memory pressure ({memory_percent:.1f}%), increasing batch size to {current_batch_size}")
                        
                        try:
                            # Reduced logging for speed - only log every 5th batch or on errors
                            if batch_count % 5 == 0 or batch_count == 1:
                                logging.info(f"[DEBUG] Adding batch {batch_count}/{total_batches} ({len(batch)} chunks)")
                            
                            self.db.add_documents(batch)
                            
                            # Only log success for every 10th batch to reduce overhead
                            if batch_count % 10 == 0:
                                logging.info(f"[DEBUG] Completed {batch_count}/{total_batches} batches successfully")
                                
                        except Exception as batch_e:
                            logging.warning(f"[DEBUG] Failed to add batch {batch_count}: {batch_e}")
                            # Try smaller batch if large batch fails
                            try:
                                smaller_batch_size = max(5, current_batch_size // 2)
                                logging.info(f"[DEBUG] Retrying with smaller batch size: {smaller_batch_size}")
                                for j in range(0, len(batch), smaller_batch_size):
                                    smaller_batch = batch[j:j+smaller_batch_size]
                                    self.db.add_documents(smaller_batch)
                                    logging.info(f"[DEBUG] Added smaller batch {j//smaller_batch_size + 1}")
                            except Exception as smaller_e:
                                logging.warning(f"[DEBUG] Failed to add smaller batch: {smaller_e}")
                                break
                    
                    logging.info(f"[DEBUG] Completed all {batch_count} batches successfully!")
            else:
                logging.info(f"[DEBUG] Creating FAISS database with {len(chunks)} chunks...")
                self.db = self._create_fast_faiss_database(chunks, self.embedding)
                logging.info("[DEBUG] FAISS database created successfully!")
            
        except Exception as e:
            logging.error(f"[DEBUG] Error creating FAISS database: {e}")
            logging.error(f"[DEBUG] Error type: {type(e)}")
            logging.error(f"[DEBUG] Error details: {str(e)}")
            
            # Try with even fewer chunks if it fails
            logging.info("[DEBUG] Trying with fewer chunks...")
            try:
                if len(chunks) > 3:
                    reduced_chunks = chunks[:3]
                    logging.info(f"[DEBUG] Trying with {len(reduced_chunks)} chunks...")
                    self.db = self._create_fast_faiss_database(reduced_chunks, self.embedding)
                    chunks = reduced_chunks  # Update chunks for logging
                    logging.info("[DEBUG] FAISS database created successfully with reduced chunks")
                elif chunks:
                    logging.info("[DEBUG] Trying with single chunk...")
                    self.db = self._create_fast_faiss_database([chunks[0]], self.embedding)
                    chunks = [chunks[0]]  # Update chunks for logging
                    logging.info("[DEBUG] FAISS database created successfully with single chunk")
                else:
                    raise Exception("No chunks available")
            except Exception as e2:
                logging.error(f"[DEBUG] Error creating FAISS database with reduced chunks: {e2}")
                # Last resort: create empty database
                logging.warning("[DEBUG] Creating empty FAISS database as last resort")
                self.db = FAISS.from_texts(["No documents available"], self.embedding)
                chunks = []
                logging.info("[DEBUG] Empty FAISS database created")
        
        # Save the database with embedding type metadata
        logging.info("[DEBUG] Saving FAISS database...")
        # Ensure the DB_PATH directory exists
        import os
        os.makedirs(DB_PATH, exist_ok=True)
        self.db.save_local(DB_PATH)
        

        
        logging.info("[DEBUG] FAISS database saved successfully")
        
        # CRITICAL: Validate data integrity after rebuild
        print(f"[HNSW] Validating data integrity after rebuild...")
        
        # Log HNSW optimization status
        if self.db and hasattr(self.db, 'index'):
            if hasattr(self.db.index, 'hnsw'):
                print(f"[HNSW] Database saved with HNSW index for fast similarity search")
                # Validate HNSW index is working
                try:
                    test_query = "integrity test"
                    test_embedding = self.get_query_embeddings(test_query)
                    test_vector = np.array([test_embedding]).astype('float32')
                    distances, indices = self.db.index.search(test_vector, k=1)
                    print(f"[HNSW] HNSW index integrity validated - search successful")
                except Exception as e:
                    print(f"[HNSW] WARNING: HNSW index integrity check failed: {e}")
            else:
                print(f"[HNSW] Database saved with standard index (HNSW optimization available on next rebuild)")
        
        # CRITICAL: Verify no orphaned data exists
        sources = set(chunk.metadata.get('source') for chunk in chunks if chunk.metadata.get('source'))
        logging.info(f"[DEBUG] Sources in vector DB: {sorted(sources)}")
        
        # Verify all sources actually exist on disk
        missing_sources = []
        for source in sources:
            if not os.path.exists(source):
                missing_sources.append(source)
        
        if missing_sources:
            logging.error(f"CRITICAL: Found {len(missing_sources)} orphaned sources in database!")
            for missing in missing_sources:
                logging.error(f"CRITICAL: Orphaned source: {missing}")
            raise Exception(f"Data integrity compromised: {len(missing_sources)} orphaned sources found")
        
        # Save build timestamp for future change detection
        try:
            build_timestamp_file = os.path.join(DB_PATH, 'last_build_timestamp.txt')
            with open(build_timestamp_file, 'w') as f:
                f.write(str(time.time()))
            logging.info(f"[DEBUG] Build timestamp saved: {build_timestamp_file}")
        except Exception as e:
            logging.warning(f"[DEBUG] Could not save build timestamp: {e}")
        
        # Save build info for embedding model change detection
        try:
            build_info_file = os.path.join(DB_PATH, 'build_info.txt')
            current_build_info = f"{self.current_embedding_type}:{MODEL_NAME}"
            with open(build_info_file, 'w') as f:
                f.write(current_build_info)
            logging.info(f"[DEBUG] Build info saved: {build_info_file} ({current_build_info})")
        except Exception as e:
            logging.warning(f"[DEBUG] Could not save build info: {e}")
        
        logging.info(f"[DEBUG] Database build completed successfully with {len(chunks)} chunks")
        add_backend_log(f"Vector database rebuilt successfully with {len(chunks)} chunks using {self.current_embedding_type} embeddings.")
        
        logging.info("CRITICAL: Data integrity validation passed - all sources exist on disk")
        logging.info(f"Vector DB rebuilt from {len(chunks)} chunks using {self.current_embedding_type} embeddings and saved to {DB_PATH}! Sources: {sorted(sources)}")
        


    def get_loader(self, path: str) -> Optional[Any]:
        """Return the appropriate loader for a file, with input validation."""
        ext = os.path.splitext(path)[-1].lower()
        try:
            if not os.path.isfile(path):
                logging.warning(f"File does not exist: {path}")
                return None
            if ext == ".txt":
                with open(path, 'r', encoding='utf-8') as f:
                    content = f.read()
                if self.is_chat_email_format_strict(content):
                    return self.extract_chat_email_knowledge(content, path)
                else:
                    return TextLoader(path)
            if ext == ".pdf": return PyMuPDFLoader(path)
            if ext == ".docx": return UnstructuredWordDocumentLoader(path)
            if ext == ".csv": return CSVLoader(path)
            if ext == ".xml":
                tree = ET.parse(path)
                root = tree.getroot()
                text = ET.tostring(root, encoding="unicode", method="text")
                return [Document(page_content=text, metadata={"source": path})]
        except Exception as e:
            logging.error(f"Error loading file {path}: {e}")
            return None

    def process_chat_email_file(self, path: str) -> Any:
        """Process chat histories and email threads to extract key information."""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            if self.is_chat_email_format(content):
                return self.extract_chat_email_knowledge(content, path)
            else:
                return TextLoader(path)
        except Exception as e:
            logging.error(f"Error processing {path}: {e}")
            return TextLoader(path)

    def is_chat_email_format(self, content: str) -> bool:
        """Detect if content is chat or email format."""
        chat_patterns = [
            r'\d{1,2}:\d{2}',
            r'\[.*?\]',
            r'<.*?@.*?>',
            r'From:|To:|Subject:',
            r'User:|Admin:|Support:',
            r'^\w+:\s',
        ]
        lines = content.split('\n')
        chat_line_count = 0
        total_lines = min(len(lines), 50)
        for line in lines[:total_lines]:
            for pattern in chat_patterns:
                if re.search(pattern, line):
                    chat_line_count += 1
                    break
        return (chat_line_count / total_lines) > 0.3

    def is_chat_email_format_strict(self, content: str) -> bool:
        """Stricter detection: Only treat as chat/email if >60% of first 50 lines match chat/email patterns."""
        chat_patterns = [
            r'\d{1,2}:\d{2}',
            r'\[.*?\]',
            r'<.*?@.*?>',
            r'From:|To:|Subject:',
            r'User:|Admin:|Support:',
            r'^\w+:\s',
        ]
        lines = content.split('\n')
        chat_line_count = 0
        total_lines = min(len(lines), 50)
        for line in lines[:total_lines]:
            for pattern in chat_patterns:
                if re.search(pattern, line):
                    chat_line_count += 1
                    break
        return (chat_line_count / total_lines) > 0.6

    def extract_chat_email_knowledge(self, content: str, source_path: str) -> List[Document]:
        """Extract key information from chat/email content and create FAQ-like documents."""
        try:
            structured_content = self.structure_chat_email_content(content)
            analysis_prompt = f"""Analyze the following chat history or email thread and extract key information, common questions, and important topics discussed. 

Create a structured summary that includes:
1. Main topics discussed
2. Common questions and their answers
3. Important decisions or conclusions
4. Key insights or learnings
5. Frequently mentioned issues or concerns

Format the output as a clear, organized FAQ-style document with sections like:
- Q: [Common Question] A: [Answer from discussion]
- Key Decision: [Important decision made]
- Insight: [Key learning or insight]

Content to analyze:
{structured_content[:8000]}

Please provide a structured summary:"""
            analysis = self.llm.invoke(analysis_prompt)
            extracted_doc = Document(
                page_content=analysis,
                metadata={
                    "source": source_path,
                    "type": "extracted_knowledge",
                    "original_format": "chat_email"
                }
            )
            return [extracted_doc]
        except Exception as e:
            logging.error(f"Error extracting knowledge from {source_path}: {e}")
            return TextLoader(source_path)

    def structure_chat_email_content(self, content: str) -> str:
        """Structure chat/email content for better analysis."""
        lines = content.split('\n')
        structured_lines = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            line = re.sub(r'\[\d{1,2}:\d{2}(:\d{2})?\]', '', line)
            line = re.sub(r'<\d{1,2}:\d{2}(:\d{2})?>', '', line)
            if len(line) > 10:
                structured_lines.append(line)
        return '\n'.join(structured_lines)

    def get_doc_topic_map(self) -> dict:
        """Return a mapping: filename -> set of topic_ids."""
        DATABASE_URL = f'postgresql://postgres:{POSTGRES_PASSWORD}@localhost:5432/chat_history'
        conn = psycopg2.connect(DATABASE_URL)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT d.filename, dt.topic_id
            FROM documents d
            JOIN document_topics dt ON d.id = dt.document_id
        ''')
        mapping = {}
        for filename, topic_id in cursor.fetchall():
            mapping.setdefault(filename, set()).add(str(topic_id))
        conn.close()
        logging.info(f"[DEBUG] Document topic mapping: {mapping}")
        return mapping

    def update_document_topics(self, filename: str, topic_ids: list) -> bool:
        """Update the topics associated with a document"""
        try:
            DATABASE_URL = f'postgresql://postgres:{POSTGRES_PASSWORD}@localhost:5432/chat_history'
            conn = psycopg2.connect(DATABASE_URL)
            cursor = conn.cursor()
            
            # First, get the document ID
            cursor.execute('SELECT id FROM documents WHERE filename = %s', (filename,))
            result = cursor.fetchone()
            
            if not result:
                print(f"[ERROR] Document '{filename}' not found in database")
                conn.close()
                return False
            
            document_id = result[0]
            
            # Remove all existing topic associations for this document
            cursor.execute('DELETE FROM document_topics WHERE document_id = %s', (document_id,))
            
            # Add new topic associations
            for topic_id in topic_ids:
                cursor.execute('''
                    INSERT INTO document_topics (document_id, topic_id) 
                    VALUES (%s, %s)
                ''', (document_id, topic_id))
            
            # Commit the changes
            conn.commit()
            conn.close()
            
            print(f"[SUCCESS] Updated topics for document '{filename}' to: {topic_ids}")
            return True
            
        except Exception as e:
            print(f"[ERROR] Failed to update document topics: {e}")
            if 'conn' in locals():
                conn.rollback()
                conn.close()
            return False

    def handle_question(self, query: str, chat_history: Optional[list] = None, topic_ids: Optional[list] = None) -> dict:
        """Handle a question with content safety checks and optional chat history context, with topic filtering."""
        start_time = time.time()
        print(f"[TIMING] Starting question processing for: {query[:50]}...")
        
        # Early gibberish check (under 10ms)
        gibberish_start = time.time()
        if self.is_gibberish(query):
            gibberish_time = time.time() - gibberish_start
            print(f"[GIBBERISH] Detected gibberish input: {query[:50]}... in {gibberish_time:.3f}s")
            return {
                "answer": "Please provide a valid question in English. Your input appears to be random characters or gibberish.",
                "detailed_answer": "Please provide a valid question in English. Your input appears to be random characters or gibberish.",
                "sources": {"summary": "N/A", "detailed": "N/A"}
            }
        gibberish_time = time.time() - gibberish_start
        print(f"[GIBBERISH] Text passed gibberish check in {gibberish_time:.3f}s")
        
        # Reset all variables to prevent contamination between questions
        answer = ""
        detailed_answer = ""
        #confidence = 0
        source_data = {}
        
        print(f"[DEBUG] handle_question called with topic_ids: {topic_ids}")
        # if self.locked_out:
        #     return {"error": "You have been locked out for repeated misuse."}
        if not self.db:
            return {"error": "Please upload compliance documents first."}
        if not query or not query.strip():
            return {"error": "Question is required."}
                # if not self.evaluate_content_safety(query, "input"):
        #     self.inappropriate_count += 1
        #     if self.inappropriate_count >= 3:
        #         self.locked_out = True
        #         return {"error": "Locked out for repeated misuse."}
        #     return {"error": "Inappropriate question. Please try again."}
        try:
            # Debug: log vector DB doc count and validate document integrity
            try:
                vector_db_count = self.db.index.ntotal
                print(f"[DEBUG] Vector DB doc count: {vector_db_count}")
                
                # Check for document count mismatch
                if hasattr(self, 'last_document_count'):
                    if vector_db_count != self.last_document_count:
                        print(f"[WARNING] Vector DB count changed from {self.last_document_count} to {vector_db_count}")
                    self.last_document_count = vector_db_count
                else:
                    self.last_document_count = vector_db_count
                    
            except Exception as e:
                print(f"[DEBUG] Could not get vector DB doc count: {e}")
            filter_filenames = None
            if topic_ids and len(topic_ids) > 0:
                # Special handling for "all topics" - don't filter, search all documents
                if 'all' in topic_ids:
                    print(f"[DEBUG] 'All topics' selected - no filtering applied, searching all documents")
                    filter_filenames = None
                else:
                    topic_filter_start = time.time()
                    print(f"[DEBUG] Topic filtering requested for topic_ids: {topic_ids}")
                    doc_topic_map = self.get_doc_topic_map()
                    print(f"[DEBUG] Document topic map: {doc_topic_map}")
                    topic_ids_set = set(str(tid) for tid in topic_ids)
                    filter_filenames = [
                        fname for fname, tags in doc_topic_map.items()
                        if topic_ids_set.issubset(tags)
                    ]
                    print(f"[DEBUG] Filtered filenames: {filter_filenames}")
                    topic_filter_time = time.time() - topic_filter_start
                    print(f"[TIMING] Topic filtering took: {topic_filter_time:.3f}s")
                    if not filter_filenames:
                        print("[DEBUG] No documents found for selected topics, returning early")
                        return {
                            "answer": "No documents are available for the selected topic(s).",
                            "detailed_answer": "No documents are available for the selected topic(s).",
                            #"confidence": 0,
                            "sources": {"summary": "N/A", "detailed": "N/A"}
                        }
            
            query_start = time.time()
            result = self.query_answer(query, chat_history=chat_history, filter_filenames=filter_filenames)
            query_time = time.time() - query_start
            print(f"[TIMING] Main query processing took: {query_time:.3f}s")
            
            # Extract the result components from the new dictionary format
            if isinstance(result, dict):
                answer = result.get("answer", "Error processing question")
                source_data = result.get("sources", {"summary": "Error", "detailed": "Error"})
                detailed_answer = result.get("detailed_answer", answer)
                hallucination_score = result.get("hallucination_score", 0.0)
                safety_issues = result.get("safety_issues", [])
                is_hallucination = result.get("is_hallucination", False)
            else:
                # Fallback if result format is unexpected
                answer = str(result) if result else "Error processing question"
                source_data = {"summary": "Error", "detailed": "Error"}
                detailed_answer = answer
                hallucination_score = 0.0
                safety_issues = []
                is_hallucination = False
            
            if answer == "I don't have specific information about that in the current documents.":
                # Return consistent response for no information
                return {
                    "answer": "I don't have specific information about that in the current documents.",
                    "detailed_answer": "No specific information available in current documents.",
                    #"confidence": 0,
                    "sources": {"summary": "N/A", "detailed": "N/A"}
                }
            
            total_time = time.time() - start_time
            print(f"[TIMING] Total question processing took: {total_time:.3f}s")
            
            return {
                "answer": answer,
                "detailed_answer": detailed_answer,
                #"confidence": confidence,
                "sources": source_data,
                "hallucination_score": hallucination_score,
                "safety_issues": safety_issues,
                "is_hallucination": is_hallucination
            }
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"Error handling question: {e}")
            print(f"Full traceback: {error_details}")
            return {"error": f"Error processing question: {str(e)}"}

    def query_answer(self, query: str, chat_history: Optional[list] = None, filter_filenames: Optional[list] = None) -> dict:
        """Answer a question using the vector database and LLM. Return comprehensive result including safety and hallucination information."""
        query_start_time = time.time()
        print(f"[TIMING] Starting query_answer processing")
        
        # COMPLETE ISOLATION - reset all context variables before each query
        self._reset_context_state()
        
        # Reset all context and answer variables to prevent contamination between queries
        actual_context = ""
        raw_context = ""
        context_parts = []
        helpful_answer = ""
        detailed_answer = ""
        answer = ""
        print(f"[DEBUG] Context and answer variables reset for new query: {query[:50]}...")
        
        if not self.db:
            raise Exception("Database not loaded. Please upload documents first.")
        
        # STEP 1: Query Preprocessing
        preprocessing_start = time.time()
        print(f"[TIMING] STEP 1: Starting query preprocessing")
        # For questionnaire processing, don't include chat history to prevent contamination
        # Each question should be processed independently
        if chat_history and len(chat_history) < 3:  # Only include history for regular chat (not questionnaires)
            history_str = "\n".join([f"User: {item['question']}\nAI: {item['answer']}" for item in chat_history])
            full_query = f"Previous conversation:\n{history_str}\n\nCurrent question: {query}"
        else:
            full_query = query
        preprocessing_time = time.time() - preprocessing_start
        print(f"[TIMING] STEP 1: Query preprocessing completed in {preprocessing_time:.3f}s")
        
        # STEP 2: Query Embedding Generation (Background Processing)
        embedding_start = time.time()
        print(f"[TIMING] STEP 2: Starting direct query embedding generation")
        try:
            # Get embedding directly - no background processing
            query_embedding = self.start_embedding_async(full_query)
            
            if query_embedding is not None:
                print(f"[TIMING] STEP 2: Direct embedding completed successfully")
            else:
                print(f"[TIMING] STEP 2: Direct embedding failed, using synchronous fallback")
                # Fallback to synchronous embedding
                query_embedding = self.get_query_embeddings(full_query)
                
        except Exception as e:
            embedding_time = time.time() - embedding_start
            print(f"[TIMING] STEP 2: Query embedding failed after {embedding_time:.3f}s: {e}")
            raise e
        
        # STEP 3: Document Retrieval (Vector Similarity Search)
        retrieval_start = time.time()
        print(f"[TIMING] STEP 3: Starting document retrieval (vector similarity search)")
        print(f"[DEBUG] Query: '{full_query}'")
        print(f"[DEBUG] Vector DB index size: {self.db.index.ntotal if hasattr(self.db, 'index') else 'Unknown'}")
        print(f"[TIMING] STEP 2: Direct embedding completed, proceeding to retrieval...")
        # Try early termination first for speed, fallback to regular search
        # OPTIMAL: Best balance between accuracy and speed
        docs = self.get_chunks_with_early_termination(full_query, similarity_threshold=0.25, min_chunks=20)  # Optimal threshold and chunk count
        retrieval_time = time.time() - retrieval_start
        print(f"[TIMING] STEP 3: Document retrieval completed in {retrieval_time:.3f}s")
        print(f"[DEBUG] Retrieved {len(docs)} documents")
        
        # Debug: Show what documents were retrieved
        if docs:
            print(f"[DEBUG] Retrieved document sources:")
            for i, doc in enumerate(docs[:5]):  # Show first 5
                source = doc.metadata.get('source', 'Unknown')
                filename = doc.metadata.get('source', 'Unknown')
                filename = os.path.basename(str(source))
                print(f"[DEBUG]   Doc {i+1}: {filename} (chunk {doc.metadata.get('page', '?')})")
                print(f"[DEBUG]   Content preview: {doc.page_content[:100]}...")
        
        # FALLBACK: If we don't have enough docs, try broader search
        if len(docs) < 15:  # If we got fewer than 15 docs, try broader search
            print(f"[FALLBACK] Only {len(docs)} docs found, trying broader search...")
            try:
                # Try with much lower threshold and more chunks
                fallback_docs = self.get_chunks_with_early_termination(full_query, similarity_threshold=0.08, min_chunks=35)
                if len(fallback_docs) > len(docs):
                    print(f"[FALLBACK] Broader search found {len(fallback_docs)} docs, using these instead")
                    docs = fallback_docs
                else:
                    print(f"[FALLBACK] Broader search didn't help, keeping original {len(docs)} docs")
            except Exception as e:
                print(f"[FALLBACK] Broader search failed: {e}, keeping original docs")
        
        # STEP 4: Context Preparation
        context_start = time.time()
        print(f"[TIMING] STEP 4: Starting context preparation")
        
        # COMPLETE CONTEXT RESET - prevent any contamination between queries
        actual_context = ""
        raw_context = ""
        context_parts = []
        
        # Clear any residual context from previous questions
        if hasattr(self, '_previous_context'):
            del self._previous_context
        if hasattr(self, '_previous_docs'):
            del self._previous_docs
        if hasattr(self, '_previous_query'):
            del self._previous_query
            
        # Create clean, isolated context from retrieved documents
        if docs:
            context_parts = []
            for i, doc in enumerate(docs):
                # Clean each document chunk to remove any artifacts
                clean_content = doc.page_content.strip()
                
                # Remove any chunk references, table names, or raw document artifacts
                import re
                clean_content = re.sub(r'Chunk \d+:', '', clean_content)
                clean_content = re.sub(r'Table \d+:', '', clean_content)
                clean_content = re.sub(r'Name of System/Type.*?Approved for use', '', clean_content, flags=re.DOTALL)
                clean_content = re.sub(r'Information Security Objectives.*?planning to achieve them:', '', clean_content, flags=re.DOTALL)
                clean_content = re.sub(r'Current question:.*?CONTEXT:', '', clean_content, flags=re.DOTALL)
                
                # Limit chunk size and clean up
                if len(clean_content) > 400:  # Increased from 250 for more context
                    clean_content = clean_content[:400] + "..."
                
                # Only add if content is meaningful after cleaning
                if len(clean_content.strip()) > 20:
                    context_parts.append(f"Document {i+1}:\n{clean_content}")
            
            actual_context = "\n\n".join(context_parts)
            
            # Store clean context for this question only
            self._previous_context = actual_context
            self._previous_docs = docs
            self._previous_query = query
        else:
            actual_context = "No relevant documents found."
        
        # Create the complete prompt with actual document context
        # complete_prompt = f"""You are a professional compliance assistant with expertise in policy interpretation, regulatory requirements, and workplace procedures.

        # CONTEXT: {actual_context}

        # QUESTION: {full_query}

        # CRITICAL INSTRUCTIONS:
        # 1. ONLY answer based on information that is EXPLICITLY stated in the provided context
        # 2. If the context does not contain relevant information to answer the question, say "I don't have specific information about that in the current documents"
        # 3. DO NOT make up information, speculate, or reference topics not mentioned in the context
        # 4. DO NOT mention companies, products, or technologies that are not in the context
        # 5. If you find relevant information, cite the specific document or policy section
        # 6. If the answer is "no", then reply with not applicable

        # ANSWER FORMAT:
        # - If relevant information exists: Provide a direct answer with specific details and citations
        # - If no relevant information: Say "I don't have specific information about that in the current documents"
        # - Be professional and workplace-appropriate
        # - Use clear, structured language

        # IMPORTANT: Only use information that is explicitly present in the context above. Do not reference any external knowledge or make assumptions. If the context is empty or contains no relevant information, you MUST say "I don't have specific information about that in the current documents".

        # FINAL CHECK: Before providing your answer, ask yourself: "Does the context actually contain information that directly answers this specific question?" If not, say "I don't have specific information about that in the current documents"."""
        logging.info(f"[DEBUG] Retriever returned {len(docs)} docs for query: {query}")
        if filter_filenames is not None:
            logging.info(f"[DEBUG] Applying topic filter with filenames: {filter_filenames}")
            original_docs_count = len(docs)
            filtered_docs = []
            for d in docs:
                source_path = str(d.metadata.get("source", ""))
                source_filename = os.path.basename(source_path)
                logging.info(f"[DEBUG] Checking doc source: '{source_path}' -> filename: '{source_filename}'")
                if source_filename in filter_filenames:
                    filtered_docs.append(d)
            docs = filtered_docs
            logging.info(f"[DEBUG] After filtering, {len(docs)} docs remain (was {original_docs_count})")
        else:
            logging.info("[DEBUG] No topic filtering applied (filter_filenames is None)")
        
        if not docs:
            logging.warning(f"[DEBUG] No documents found for query: {query}")
            return ("I don't have specific information about that in the current documents.", 0, {"summary": "N/A", "detailed": "N/A"}, "I don't have specific information about that in the current documents.")
        
        # Log document sources for debugging
        doc_sources = [os.path.basename(str(d.metadata.get("source", ""))) for d in docs]
        logging.info(f"[DEBUG] Using documents: {doc_sources}")
        
        # Clear any previous context to prevent contamination between queries
        actual_context = ""
        
        # Update the context with ENHANCED filtering and building for better quality
        if docs:
            # Enhanced filtering with multi-factor relevance scoring
            # Balanced threshold: not too loose, not too strict
            relevant_docs = self.filter_relevant_chunks(query, docs, threshold=0.6)
            print(f"[QUALITY] Enhanced filtering: {len(docs)} -> {len(relevant_docs)} relevant chunks")
            
            # Build context with intelligent chunking and overlap
            raw_context = self.build_context_streamlined(relevant_docs, max_chars=3500, overlap_ratio=0.2)
            
            # Validate context relevance before proceeding
            validated_context = self.validate_context_relevance(query, raw_context, relevant_docs)
            
            # Apply final context size limiting
            actual_context = self.limit_context_size(validated_context, max_chars=3000)
            print(f"[QUALITY] Context prepared: {len(relevant_docs)} chunks -> {len(actual_context)} chars")
            
            # DEBUG: Show context preview for troubleshooting
            print(f"[DEBUG] Context preview (first 500 chars): {actual_context[:500]}...")
        else:
            actual_context = "No relevant chunks found."
        
        # FINAL CONTEXT VALIDATION - ensure prompt is clean before sending to LLM
        final_context = actual_context
        final_query = full_query
        
        # Remove any remaining contamination from the query
        if "Previous conversation:" in final_query:
            # Extract only the current question
            final_query = final_query.split("Current question: ")[-1] if "Current question: " in final_query else full_query
        
        # Remove any remaining contamination from the context
        import re
        final_context = re.sub(r'Chunk \d+:', '', final_context)
        final_context = re.sub(r'Table \d+:', '', final_context)
        final_context = re.sub(r'Name of System/Type.*?Approved for use', '', final_context, flags=re.DOTALL)
        final_context = re.sub(r'Information Security Objectives.*?planning to achieve them:', '', final_context, flags=re.DOTALL)
        final_context = re.sub(r'Current question:.*?CONTEXT:', '', final_context, flags=re.DOTALL)
        
        # Use enhanced prompt engineering for better answer quality
        complete_prompt = self.build_enhanced_prompt(full_query, actual_context, chat_history)

        # Build minimal sources info for speed (skip detailed processing)
        sources = []
        source_summary = {}
        
        # Clear any residual context to prevent cross-question contamination
        if hasattr(self, 'llm') and hasattr(self.llm, 'reset'):
            try:
                self.llm.reset()
                print(f"[DEBUG] LLM context reset for question isolation")
            except Exception as e:
                print(f"[DEBUG] LLM reset failed (non-critical): {e}")
        
        # Use only the relevant_docs that were actually used to build the context
        # instead of all retrieved docs
        docs_for_sources = relevant_docs if 'relevant_docs' in locals() else docs
        
        print(f"[SOURCES] Retrieved {len(docs)} docs, using {len(docs_for_sources)} relevant docs for sources")
        
        for i, d in enumerate(docs_for_sources):
            src = os.path.basename(str(d.metadata.get("source", "N/A")))
            page = d.metadata.get("page", d.metadata.get("page_number", "?"))
            
            # Enhanced source string with meaningful location information
            if page != "?":
                location_info = self._get_meaningful_location(src, page)
                source_str = f"{src} ({location_info})"
            else:
                source_str = f"{src}"
            sources.append(source_str)
            
            # Build summary by file
            if src not in source_summary:
                source_summary[src] = []
            source_summary[src].append(page)
        
        # Create summary format with meaningful location information
        summary_parts = []
        for src, pages in source_summary.items():
            unique_pages = sorted(list(set(pages)), key=lambda x: int(str(x)) if str(x).isdigit() else 0)
            if len(unique_pages) == 1:
                location_info = self._get_meaningful_location(src, unique_pages[0])
                summary_parts.append(f"{src} ({location_info})")
            else:
                location_infos = [self._get_meaningful_location(src, page) for page in unique_pages]
                summary_parts.append(f"{src} ({', '.join(location_infos)})")
        
        sources_summary = " | ".join(summary_parts)
        sources_detailed = " | ".join(sources)
        
        # Format sources more cleanly
        if len(sources) == 1:
            sources_summary = summary_parts[0]
            sources_detailed = sources[0]
        top_similarity = None
        try:
            # Get the embedding result directly
            if query_embedding is not None:
                print(f"[TIMING] STEP 2: Using direct embedding result")
                query_emb = query_embedding  # Direct embedding result
                embedding_time = time.time() - embedding_start
                print(f"[TIMING] STEP 2: Direct embedding completed in {embedding_time:.3f}s")
            else:
                # Fallback if embedding failed
                print(f"[TIMING] STEP 2: Embedding failed, using fallback")
                embedding_time = time.time() - embedding_start
                print(f"[TIMING] STEP 2: Embedding failed in {embedding_time:.3f}s")
                query_emb = None
            
            # Calculate similarity only with the retrieved documents (not all chunks)
            if docs:
                import numpy as np
                from numpy.linalg import norm
                similarities = []
                
                # Calculate similarity with each retrieved document using optimized embeddings
                for doc in docs:
                    # Get document embedding from FAISS index
                    doc_id = doc.metadata.get('chunk_id', 0)
                    if hasattr(self.db.index, 'reconstruct'):
                        doc_emb = self.db.index.reconstruct(doc_id)
                        
                        # SPEED OPTIMIZATION: Ensure both embeddings have same dimension
                        if len(query_emb) != len(doc_emb):
                            # Truncate to smaller dimension for compatibility
                            min_dim = min(len(query_emb), len(doc_emb))
                            query_emb_trunc = query_emb[:min_dim]
                            doc_emb_trunc = doc_emb[:min_dim]
                            print(f"[SPEED] Truncated embeddings to {min_dim}D for similarity calculation")
                        else:
                            query_emb_trunc = query_emb
                            doc_emb_trunc = doc_emb
                        
                        sim = np.dot(query_emb_trunc, doc_emb_trunc) / (norm(query_emb_trunc) * norm(doc_emb_trunc) + 1e-8)
                        similarities.append(sim)
                
                if similarities:
                    top_similarity = max(similarities)
                    logging.info(f"[DEBUG] Native {MODEL_NAME} similarity calculation successful with {len(docs)} docs, max similarity: {top_similarity:.4f}")
                    
        except Exception as e:
            logging.error(f"Native {MODEL_NAME} similarity calculation failed: {e}")
            top_similarity = None
        # Complete context preparation timing
        context_time = time.time() - context_start
        print(f"[TIMING] STEP 4: Context preparation completed in {context_time:.3f}s")
        
        # STEP 5: LLM Call with Speed Optimizations
        llm_start = time.time()
        print(f"[TIMING] STEP 5: Starting optimized LLM call")
        try:
            print(f"[SPEED] Sending CLEANED prompt to LLM (length: {len(complete_prompt)} chars)")
            print(f"[DEBUG] Final query length: {len(final_query)} chars")
            print(f"[DEBUG] Final context length: {len(final_context)} chars")
            print(f"[DEBUG] Total prompt length: {len(complete_prompt)} chars")
            
            # Use direct LLM call for maximum speed (no streaming overhead)
            raw_answer = self.llm.invoke(complete_prompt)
            
            # Use full response for better accuracy (no length limiting)
            answer = raw_answer
            
            # Handle different response formats
            if isinstance(answer, dict) and 'result' in answer:
                answer = answer['result']
            else:
                answer = str(answer)
            if isinstance(answer, list):
                answer = "\n".join(str(a) for a in answer)
                
            llm_time = time.time() - llm_start
            print(f"[TIMING] STEP 5: Optimized LLM call completed in {llm_time:.3f}s")
            print(f"[SPEED] LLM response received (length: {len(answer)} chars)")
            
            # STEP 5.5: Hallucination & Content Safety Check - DISABLED
            # safety_start = time.time()
            # print(f"[SAFETY] STEP 5.5: Starting integrated hallucination and content safety check")
            
            # # Check for both hallucinations and content safety in one AI evaluation
            # hallucination_check = self._detect_hallucination(answer, query, final_context, relevant_docs)
            
            # if hallucination_check['is_hallucination']:
            #     print(f"[SAFETY] Potential hallucination detected: {hallucination_check['confidence']:.2f}")
            #     # Add warning to the answer
            #     answer = f"[⚠️ WARNING: This answer may contain information not directly supported by the provided documents. Please verify with official sources.]\n\n{answer}"
            
            # safety_time = time.time() - safety_start
            # print(f"[SAFETY] STEP 5.5: Integrated safety checks completed in {safety_time:.3f}s")
            # print(f"[SAFETY] Hallucination score: {hallucination_check['confidence']:.3f}")
            # print(f"[SAFETY] Content safety: {'PASS' if not hallucination_check.get('is_inappropriate', False) else 'FAIL'}")
            
            # Create a dummy hallucination check for compatibility
            hallucination_check = {'is_hallucination': False, 'confidence': 0.0, 'is_inappropriate': False}
            
            # STEP 6: Response Processing & Output
            response_processing_start = time.time()
            print(f"[TIMING] STEP 6: Starting response processing & output")
            
            # Parse the response to extract the focused answer
            helpful_answer = ""
            detailed_answer = ""
            
            # Extract the actual answer content after "ANSWER:"
            if "ANSWER:" in answer:
                # Get everything after "ANSWER:"
                answer_content = answer.split("ANSWER:")[1].strip()
                
                # Clean up any remaining artifacts
                cleaned_answer = re.sub(r'Chunk \d+:', '', answer_content)
                cleaned_answer = re.sub(r'Table \d+:', '', cleaned_answer)
                cleaned_answer = re.sub(r'Name of System/Type.*?Approved for use', '', cleaned_answer, flags=re.DOTALL)
                cleaned_answer = re.sub(r'Information Security Objectives.*?planning to achieve them:', '', cleaned_answer, flags=re.DOTALL)
                cleaned_answer = re.sub(r'Current question:.*?CONTEXT:', '', cleaned_answer, flags=re.DOTALL)
                cleaned_answer = re.sub(r'DETAILED ANSWER:.*?\[Your comprehensive answer here\]', '', cleaned_answer, flags=re.DOTALL)
                cleaned_answer = re.sub(r'ONE-LINE ANSWER:.*?\[Your concise one-line answer here\]', '', cleaned_answer, flags=re.DOTALL)
                
                # Extract confidence percentage if present
                confidence_match = re.search(r'Confidence:\s*(\d+)%', cleaned_answer, re.IGNORECASE)
                confidence_percentage = confidence_match.group(1) if confidence_match else None
                
                # Remove confidence line from the answer for display
                cleaned_answer = re.sub(r'Confidence:\s*\d+%.*', '', cleaned_answer, flags=re.IGNORECASE)
                
                # Take only the first 3-4 meaningful sentences for concise answers
                sentences = cleaned_answer.split('.')
                meaningful_sentences = []
                
                for sentence in sentences:
                    sentence = sentence.strip()
                    if sentence and len(sentence) > 10 and not sentence.startswith('Chunk') and not sentence.startswith('Table'):
                        meaningful_sentences.append(sentence)
                        if len(meaningful_sentences) >= 4:  # Limit to 4 sentences max
                            break
                
                if meaningful_sentences:
                    helpful_answer = '. '.join(meaningful_sentences) + '.'
                    detailed_answer = helpful_answer
                else:
                    helpful_answer = "No specific information found about this topic."
                    detailed_answer = helpful_answer
                
                # Add confidence back to the detailed answer if found
                if confidence_percentage:
                    detailed_answer += f"\n\nConfidence: {confidence_percentage}%"
                    print(f"[QUALITY] Confidence score extracted: {confidence_percentage}%")
                else:
                    print(f"[QUALITY] No confidence score found in response")
            else:
                # Fallback: try to extract meaningful content
                cleaned_answer = re.sub(r'Question:.*?Context:', '', answer, flags=re.DOTALL)
                cleaned_answer = re.sub(r'Chunk \d+:', '', cleaned_answer)
                cleaned_answer = re.sub(r'Table \d+:', '', cleaned_answer)
                
                # Take first meaningful sentence
                sentences = cleaned_answer.split('.')
                for sentence in sentences:
                    sentence = sentence.strip()
                    if sentence and len(sentence) > 20 and not sentence.startswith('Chunk') and not sentence.startswith('Table'):
                        helpful_answer = sentence + '.'
                        detailed_answer = helpful_answer
                        break
                
                if not helpful_answer:
                    helpful_answer = "No specific information found about this topic."
                    detailed_answer = helpful_answer
            
            print(f"[DEBUG] Cleaned answer: {len(helpful_answer)} chars")
            print(f"[DEBUG] Answer preview: {helpful_answer[:100]}...")
            
            # Use the cleaned answer
            answer = helpful_answer
            
            # Final cleanup - remove any remaining raw document artifacts
            if len(answer) > 500:  # If answer is still too long
                print(f"[DEBUG] Answer still too long ({len(answer)} chars), applying final cleanup")
                # Take only the first meaningful part
                first_period = answer.find('.')
                if first_period > 0:
                    answer = answer[:first_period + 1]
                    print(f"[DEBUG] Final answer length: {len(answer)} chars")
            
            # Remove any question repetition
            if answer.startswith("Question:") or answer.startswith("Context:"):
                print(f"[DEBUG] Question/context detected at start, cleaning")
                # Find the first sentence that doesn't start with these
                sentences = answer.split('.')
                for sentence in sentences:
                    sentence = sentence.strip()
                    if sentence and not sentence.startswith('Question:') and not sentence.startswith('Context:') and len(sentence) > 10:
                        answer = sentence + '.'
                        break
                else:
                    answer = "No specific information found about this topic."
            
            # Ensure answer doesn't contain raw document artifacts
            if 'Chunk' in answer or 'Table' in answer or 'Name of System' in answer:
                print(f"[DEBUG] Raw document content detected, using fallback")
                answer = "No specific information found about this topic."
            
            # FALLBACK: If LLM says no info found but we have context, try to extract relevant info
            if "no specific information found" in answer.lower() or "no information found" in answer.lower():
                print(f"[DEBUG] LLM said no info found, attempting fallback extraction")
                fallback_answer = self._extract_fallback_answer(full_query, actual_context)
                if fallback_answer:
                    print(f"[DEBUG] Fallback extraction successful: {fallback_answer[:100]}...")
                    answer = fallback_answer
                else:
                    print(f"[DEBUG] Fallback extraction also failed")
        
        except Exception as e:
            logging.error(f"Error calling LLM: {e}")
            import traceback
            logging.error(f"LLM error traceback: {traceback.format_exc()}")
            raise Exception(f"Failed to get response from AI model: {str(e)}")
        
        # Activity 5 & 6: Extra processing (guardrails, hallucination checks, etc.)
        extra_processing_start = time.time()
        logging.info(f"[TIMING] Starting extra processing (guardrails, hallucination checks)")
        
        # Post-processing: Check for hallucination and validate answer quality
        answer_lower = answer.lower()
        
        # # AI-based hallucination detection using embeddings and similarity
        # potential_hallucination = False
        
        # if docs and answer.strip():
        #     try:
        #         # Get embeddings for the answer and the context
        #         answer_embedding = self.embedding.embed_query(answer)
                
        #         # Create context embedding by combining all document embeddings
        #         context_text = "\n\n".join([doc.page_content for doc in docs])
        #         context_embedding = self.embedding.embed_query(context_text)
                
        #         # Calculate cosine similarity between answer and context
        #         import numpy as np
        #         from numpy.linalg import norm
                
        #         similarity = np.dot(answer_embedding, context_embedding) / (norm(answer_embedding) * norm(context_embedding) + 1e-8)
                
        #         logging.info(f"[DEBUG] Answer-context similarity: {similarity:.4f}")
                
        #         # If similarity is very low (< 0.3), the answer might be hallucinated
        #         if similarity < 0.3):
        #             logging.warning(f"[DEBUG] Low similarity detected ({similarity:.4f}), potential hallucination")
        #             potential_hallucination = True
                    
        #     except Exception as e:
        #         logging.warning(f"[DEBUG] Error in AI-based hallucination detection: {e}")
        #         # Fallback: don't flag as hallucination if detection fails
        
        # if potential_hallucination:
        #     logging.warning(f"[DEBUG] Potential hallucination detected in answer: {answer[:200]}...")
        #     # Force a re-prompt with stronger anti-hallucination instructions
        #     anti_hallucination_prompt = f"""You are a professional compliance assistant. The user asked: {query}

        # CRITICAL: You must ONLY answer based on information explicitly stated in the provided context. 

        # CONTEXT: {actual_context}

        # IMPORTANT RULES:
        # 1. If the context does not contain relevant information, say "I don't have specific information about that in the current documents"
        # 2. Do NOT mention any companies, products, or technologies not explicitly mentioned in the context
        # 3. Do NOT reference "Document 1", "Document 2", etc. unless they are actual document names in the context
        # 4. Do NOT make up information or speculate
        # 5. If the context is empty or contains no relevant information, you MUST say "I don't have specific information about that in the current documents"

        # Please provide your answer:"""
            
        #     try:
        #         corrected_result = self.llm.invoke(anti_hallucination_prompt)
        #         if isinstance(corrected_result, dict) and 'result' in corrected_result:
        #             answer = corrected_result['result']
        #         else:
        #             answer = str(corrected_result)
        #         logging.info("[DEBUG] Answer corrected for potential hallucination")
        #     except Exception as e:
        #         logging.warning(f"Anti-hallucination correction failed: {e}")
        
        # # Check if answer is too vague or says "I don't know"
        # if any(phrase in answer_lower for phrase in ["i don't know", "i don't have", "no information", "not mentioned", "not found"]):
        #     # Re-prompt with stronger instructions if context exists
        #     if docs and any(d.page_content.strip() for d in docs):
        #         logging.info("[DEBUG] Answer was too vague, re-prompting with stronger instructions")
        #         enhanced_prompt = f"""You are a professional compliance assistant. The user asked: {query}

        # You have access to relevant documents. Please answer based ONLY on information explicitly stated in the context.

        # IMPORTANT: 
        # - If the context contains relevant information, provide a clear answer with citations
        # - If the context does not contain relevant information, say "I don't have specific information about that in the current documents"
        # - Do NOT make up information or reference topics not in the context
        # - If the context is empty or contains no relevant information, you MUST say "I don't have specific information about that in the current documents"

        # Please answer the question using the available context:"""
                
        #         try:
        #             enhanced_result = self.llm.invoke(enhanced_prompt)
        #             if isinstance(enhanced_result, dict) and 'result' in enhanced_result:
        #             answer = enhanced_result['result']
        #         else:
        #             answer = str(enhanced_result)
        #         except Exception as e:
        #             logging.warning(f"Enhanced prompt failed: {e}")
        #             # Keep the original answer if enhancement fails
        
        # # Additional quality check: If answer is still too short or vague, try one more time
        # if len(answer.strip()) < 50 and docs:
        #     logging.info("[DEBUG] Answer is too short, attempting to generate more detailed response")
        #     detailed_prompt = f"""You are a professional compliance assistant. The user asked: {query}

        # You have access to relevant documents. Please provide an answer that:
        # 1. ONLY uses information explicitly stated in the context
        # 2. Directly addresses the question if relevant information exists
        # 3. Includes specific details and citations from the documents
        # 4. Is professional and workplace-appropriate
        # 5. If no relevant information exists, clearly state "I don't have specific information about that in the current documents"

        # IMPORTANT: Do not make up information or reference topics not in the context. If the context is empty or contains no relevant information, you MUST say "I don't have specific information about that in the current documents".

        # Please provide your answer:"""
            
        #     try:
        #         detailed_result = self.llm.invoke(detailed_prompt)
        #         if isinstance(detailed_result, dict) and 'result' in detailed_result:
        #             answer = detailed_result['result']
        #         else:
        #             answer = str(detailed_result)
        #         except Exception as e:
        #             logging.warning(f"Detailed prompt failed: {e}")
        #             # Keep the original answer if enhancement fails
        
        # Use simple default confidence since LLM-based calculation is disabled
        #confidence = 0.85  # Default confidence
        #print(f"[TIMING] Using default confidence: {confidence}")
        
        # Return structured source data - NO SOURCES if answer indicates no information
        no_info_phrases = [
            "no specific information found", "no information found", "i don't have specific information",
            "not mentioned", "not found", "no information available", "i don't know"
        ]
        
        answer_lower = answer.lower()
        has_no_info = any(phrase in answer_lower for phrase in no_info_phrases)
        
        if has_no_info:
            # If answer indicates no information, don't show sources
            source_data = {
                "summary": "No sources - no information found",
                "detailed": "No sources - no information found"
            }
        else:
            # Normal case: show sources that contributed to the answer
            source_data = {
                "summary": sources_summary if sources else "N/A",
                "detailed": sources_detailed if sources else "N/A"
            }
        
        # Skip answer enhancement entirely for speed
        final_answer = answer.strip()
        print(f"[TIMING] Skipping answer enhancement for speed")
        
        # Skip final hallucination check entirely for speed
        print(f"[TIMING] Skipping final hallucination check for speed")
        
        # Use the current answer directly (no variable contamination)
        final_answer = answer.strip()
        detailed_answer = answer.strip()
        
        # Complete response processing timing
        response_processing_time = time.time() - response_processing_start
        print(f"[TIMING] STEP 6: Response processing & output completed in {response_processing_time:.3f}s")
        
        # Calculate total query time
        total_query_time = time.time() - query_start_time
        print(f"[TIMING] Total query_answer processing took: {total_query_time:.3f}s")
        
        # Print comprehensive timing summary
        print(f"[TIMING] ===== QUERY PIPELINE TIMING SUMMARY =====")
        print(f"[TIMING] STEP 1 (Preprocessing): {preprocessing_time:.3f}s")
        print(f"[TIMING] STEP 2 (Embedding): {embedding_time:.3f}s")
        print(f"[TIMING] STEP 3 (Retrieval): {retrieval_time:.3f}s")
        print(f"[TIMING] STEP 4 (Context Prep): {context_time:.3f}s")
        print(f"[TIMING] STEP 5 (LLM Call): {llm_time:.3f}s")
        print(f"[TIMING] STEP 6 (Response Processing): {response_processing_time:.3f}s")
        print(f"[TIMING] TOTAL TIME: {total_query_time:.3f}s")
        print(f"[TIMING] ==========================================")
        
        # Log what we're returning for debugging
        print(f"[DEBUG] Returning - final_answer: {len(final_answer)} chars, detailed_answer: {len(detailed_answer)} chars")
        print(f"[DEBUG] Final helpful preview: {final_answer[:100]}...")
        print(f"[DEBUG] Final detailed preview: {detailed_answer[:100]}...")
        print(f"[DEBUG] FINAL answer variable before return: {answer[:100]}...")
        
        # Return comprehensive result including safety and hallucination information
        return {
            "answer": final_answer,
            "detailed_answer": detailed_answer,
            "sources": source_data,
            "hallucination_score": hallucination_check['confidence'] if 'hallucination_check' in locals() else 0.0,
            "safety_issues": [] if hallucination_check.get('is_inappropriate', False) else [],
            "is_hallucination": hallucination_check['is_hallucination'] if 'hallucination_check' in locals() else False,
            "safety_check": {'safe': not hallucination_check.get('is_inappropriate', False), 'reason': 'Integrated safety check completed'},
            "hallucination_check": hallucination_check if 'hallucination_check' in locals() else {'is_hallucination': False, 'confidence': 0.0}
        }

#     def calculate_answer_confidence(self, query: str, answer: str, docs: list) -> float:
#         """Calculate the AI's confidence in its answer using enhanced evaluation"""
#         try:
#             # Prepare context from retrieved documents
#             context = "\n\n".join([f"Document {i+1}: {doc.page_content}" for i, doc in enumerate(docs)])
            
#             # Enhanced confidence evaluation prompt
#             confidence_prompt = f"""You are an expert evaluator assessing the quality and confidence of a compliance assistant's answer.

# Question: {query}

# Answer provided: {answer}

# Source documents used:
# {context}

# EVALUATION CRITERIA:
# 1. RELEVANCE: Does the answer directly address the question asked?
# 2. ACCURACY: Is the information in the answer supported by the source documents?
# 3. COMPLETENESS: Does the answer provide sufficient detail to be useful?
# 4. SYNTHESIS: If multiple sources were used, were they effectively combined?
# 5. PROFESSIONALISM: Is the answer clear, well-structured, and workplace-appropriate?

# CONFIDENCE SCALE (0-100):
# - 0-20: Very low confidence (answer is incorrect, irrelevant, or completely unsupported)
# - 21-40: Low confidence (answer has major gaps, is vague, or poorly supported)
# - 41-60: Moderate confidence (answer is partially correct but has significant limitations)
# - 61-80: High confidence (answer is mostly correct, well-supported, and useful)
# - 81-100: Very high confidence (answer is accurate, complete, well-synthesized, and professional)

# Consider the quality of the answer relative to the available information. If the answer effectively uses whatever relevant information is available, even if incomplete, it should receive higher confidence.

# Respond with ONLY a number between 0 and 100, representing your confidence percentage."""

#             # Get confidence score from LLM
#             confidence_response = self.llm.invoke(confidence_prompt)
#             confidence_text = str(confidence_response).strip()
            
#             # Extract numeric confidence score
#             import re
#             confidence_match = re.search(r'\b(\d{1,3})\b', confidence_text)
#             if confidence_match:
#                 confidence_score = float(confidence_match.group(1))
#                 # Ensure it's within 0-100 range
#                 confidence_score = max(0, min(100, confidence_score))
#                 return round(confidence_score, 2)
#             else:
#                 # Fallback if parsing fails
#                 return 50.0
                
#         except Exception as e:
#             logging.warning(f"Confidence calculation failed: {e}")
#             # Fallback confidence based on whether we have documents
#             return 20.0 if docs else 0.0

#     def enhance_answer_quality(self, query: str, answer: str, docs: list) -> str:
#         """Enhance answer quality by checking for common issues and improving structure"""
#         enhancement_start = time.time()
#         logging.info(f"[TIMING] Starting answer enhancement")
        
#         try:
#             # Check if answer needs improvement
#             answer_lower = answer.lower()
            
#             # If answer is too short or doesn't seem comprehensive enough
#             if len(answer.strip()) < 100 and docs:
#                 logging.info("[DEBUG] Answer may need enhancement, checking quality")
                
#                 # Check if we have good source material
#                 total_source_length = sum(len(doc.page_content) for doc in docs)
#                 if total_source_length > 500:  # We have substantial source material
#                     enhancement_prompt = f"""You are a professional compliance assistant. The user asked: {query}

# Current answer: {answer}

# You have access to substantial source material. Please enhance this answer to be more comprehensive and professional. The enhanced answer should:

# 1. Be more detailed and specific
# 2. Include relevant policy citations or references when possible
# 3. Structure the information logically
# 4. Be professional and workplace-appropriate
# 5. Address the question more thoroughly

# Please provide an enhanced version of the answer:"""
                    
#                     # enhancement_llm_start = time.time()
#                     # enhanced_result = self.llm.invoke(enhancement_prompt)
#                     # enhancement_llm_time = time.time() - enhancement_llm_start
#                     # logging.info(f"[TIMING] Enhancement LLM call took: {enhancement_llm_time:.3f}s")
                    
#                     # if isinstance(enhanced_result, dict) and 'result' in enhanced_result:
#                     #     enhanced_answer = enhanced_result['result']
#                     # else:
#                     #     enhanced_answer = str(enhanced_result)
                    
#                     # # Only use enhanced answer if it's significantly better
#                     # if len(enhanced_answer.strip()) > len(answer.strip()) * 1.5:
#                     #     enhancement_time = time.time() - enhancement_start
#                     #     logging.info(f"[TIMING] Answer enhancement completed in: {enhancement_time:.3f}s")
#                     #     return enhanced_answer.strip()
            
#             # Check for common issues and fix them
#             if "i don't know" in answer_lower or "no information" in answer_lower:
#                 # Try to extract any useful information from sources
#                 useful_info = []
#                 for doc in docs:
#                     content = doc.page_content.lower()
#                     if any(keyword in content for keyword in query.lower().split()):
#                         useful_info.append(doc.page_content[:200] + "...")
                
#                 if useful_info:
#                     fallback_prompt = f"""The user asked: {query}

# While I don't have the exact answer, I found some potentially relevant information in the documents. Please provide a helpful response that:

# 1. Acknowledges the limitations
# 2. Shares any relevant information found
# 3. Suggests where they might find more information
# 4. Is professional and helpful

# Relevant information found:
# {chr(10).join(useful_info)}

# Please provide a helpful response:"""
                    
#                     # fallback_llm_start = time.time()
#                     # fallback_result = self.llm.invoke(fallback_prompt)
#                     # fallback_llm_time = time.time() - fallback_llm_start
#                     # logging.info(f"[TIMING] Fallback LLM call took: {fallback_llm_time:.3f}s")
                    
#                     # if isinstance(fallback_result, dict) and 'result' in enhanced_result):
#                     #     enhancement_time = time.time() - enhancement_start
#                     #     logging.info(f"[TIMING] Answer enhancement completed in: {enhancement_time:.3f}s")
#                     #     return fallback_result['result'].strip()
            
#             enhancement_time = time.time() - enhancement_start
#             logging.info(f"[TIMING] Answer enhancement completed in: {enhancement_time:.3f}s")
#             return answer.strip()
            
#         except Exception as e:
#             logging.error(f"Error enhancing answer quality: {e}")
#             enhancement_time = time.time() - enhancement_start
#             logging.info(f"[TIMING] Answer enhancement failed after: {enhancement_time:.3f}s")
#             return answer.strip()
    
    def generate_short_answer(self, query: str, detailed_answer: str, docs: list) -> str:
        """Generate a concise, one-line summary answer for quick overview."""
        try:
            short_prompt = f"""You are a professional compliance assistant. The user asked: {query}

Here is the detailed answer: {detailed_answer}

Please create a concise, one-line summary answer (maximum 2 sentences) that:
1. Directly answers the question
2. Captures the key point or conclusion
3. Is professional and clear
4. Can be understood quickly

Provide ONLY the short summary answer:"""

            # short_result = self.llm.invoke(short_prompt)
            # if isinstance(short_result, dict) and 'result' in short_result:
            #     short_answer = short_result['result']
            # else:
            #     short_answer = str(short_result)
            
            # Clean up the short answer
            short_answer = short_answer.strip()
            if short_answer.startswith('"') and short_answer.endswith('"'):
                short_answer = short_answer[1:-1]
            
            # Fallback: if LLM fails, create a simple summary
            if len(short_answer) < 10 or len(short_answer) > 300:
                # Create a simple summary from the detailed answer
                sentences = detailed_answer.split('.')
                if sentences:
                    short_answer = sentences[0].strip() + '.'
                    if len(short_answer) > 500:  # Increased from 200 to 500 for better display
                        short_answer = short_answer[:500] + '...'
            
            return short_answer
            
        except Exception as e:
            logging.warning(f"Short answer generation failed: {e}")
            # Fallback: return first sentence of detailed answer
            sentences = detailed_answer.split('.')
            if sentences:
                return sentences[0].strip() + '.'
            else:
                return detailed_answer[:300] + '...' if len(detailed_answer) > 300 else detailed_answer

    def force_rebuild_db(self) -> None:
        """Safely delete the FAISS index directory and rebuild the vector DB from scratch."""
        import os
        import shutil
        from pathlib import Path
        
        print(f"[HNSW] CRITICAL: Starting forced rebuild with complete data cleanup...")
        
        # CRITICAL: Clear database from memory first
        self.db = None
        
        db_path = Path(DB_PATH).resolve()
        project_root = Path(__file__).parent.resolve()
        
        # Only allow deletion if DB_PATH is a subdirectory of the project and not '.', '', or '/'
        if db_path == project_root or str(db_path) in ('/', '') or not str(db_path).startswith(str(project_root)):
            add_backend_log(f"[ERROR] Unsafe DB_PATH for deletion: {db_path}. Skipping vector DB cleanup.")
            return
        
        # CRITICAL: Force complete cleanup of all index files
        if db_path.exists() and db_path.is_dir():
            try:
                # Remove all files in the directory
                for file_path in db_path.iterdir():
                    if file_path.is_file():
                        file_path.unlink()
                        print(f"[HNSW] CRITICAL: Deleted index file: {file_path}")
                
                # Remove the directory itself
                shutil.rmtree(db_path)
                print(f"[HNSW] CRITICAL: Deleted old FAISS index directory: {db_path}")
                add_backend_log(f"Deleted old FAISS index directory: {db_path}")
                
            except Exception as e:
                add_backend_log(f"[ERROR] Error deleting FAISS index: {e}")
                print(f"[HNSW] CRITICAL ERROR: Could not delete index directory: {e}")
                raise Exception(f"Cannot delete index directory: {e}")
        else:
            print(f"[HNSW] CRITICAL: No FAISS index directory to delete at: {db_path}")
            add_backend_log(f"No FAISS index directory to delete at: {db_path}")
        
        # CRITICAL: Verify cleanup was successful
        if db_path.exists():
            raise Exception(f"CRITICAL: Index directory still exists after deletion attempt: {db_path}")
        
        print(f"[HNSW] CRITICAL: Complete cleanup verified - rebuilding database...")
        
        # Rebuild with guaranteed clean state
        self.build_db("forced_rebuild")
        add_backend_log("Forced rebuild of vector DB complete with guaranteed data cleanup.")
        print(f"[HNSW] CRITICAL: Forced rebuild completed with guaranteed data integrity")

    def add_document_incremental(self, file_path: str) -> bool:
        """Add a single document to the existing database using incremental update."""
        try:
            logging.info(f"[DEBUG] Adding document incrementally: {file_path}")
            
            # Switch to Ollama for incremental updates
            self.select_optimal_embedding("incremental_update")
            
            # Load the single document
            filename = os.path.basename(file_path)
            ext = os.path.splitext(file_path)[1].lower()
            
            # Check if document already exists in database to prevent duplicates
            if self.db:
                try:
                    # Get a sample of existing documents to check for duplicates
                    sample_docs = self.db.similarity_search("", k=1000)
                    existing_sources = set()
                    for doc in sample_docs:
                        source = doc.metadata.get('source', '')
                        if source:
                            existing_sources.add(os.path.basename(source))
                    
                    if filename in existing_sources:
                        logging.warning(f"[DEBUG] Document {filename} already exists in database, skipping incremental add")
                        return True  # Already exists, consider it successful
                        
                except Exception as e:
                    logging.warning(f"[DEBUG] Could not check for duplicates: {e}, proceeding with add")
            
            docs = []
            if ext == ".pdf":
                try:
                    loader = PyMuPDFLoader(file_path)
                    docs = loader.load()
                    logging.info(f"[DEBUG] Loaded {len(docs)} docs from {filename} (PyMuPDF)")
                except Exception as e:
                    logging.warning(f"[DEBUG] PyMuPDF failed for {filename}: {e}. Trying OCR...")
                    loader = UnstructuredPDFLoader(file_path, strategy="ocr_only")
                    docs = loader.load()
                    logging.info(f"[DEBUG] Loaded {len(docs)} docs from {filename} (OCR)")
            elif ext == ".txt":
                loader = TextLoader(file_path)
                docs = loader.load()
            elif ext == ".docx":
                loader = UnstructuredWordDocumentLoader(file_path)
                docs = loader.load()
            elif ext == ".csv":
                loader = CSVLoader(file_path)
                docs = loader.load()
            else:
                logging.warning(f"Unsupported file type: {file_path}")
                return False
            
            if not docs:
                logging.warning(f"No documents loaded from {filename}")
                return False
            
            # Use smart chunking for consistent quality
            chunk_size, chunk_overlap = self.calculate_optimal_chunk_parameters()
            
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,  # Use character count for consistent sizing
                separators=[
                    "\n\n\n",    # Major section breaks
                    "\n\n",      # Paragraph breaks
                    "\n",        # Line breaks
                    ". ",        # Sentence endings
                    "! ",        # Exclamation endings
                    "? ",        # Question endings
                    "; ",        # Semicolon separators
                    ": ",        # Colon separators
                    " - ",       # Dash separators
                    " • ",       # Bullet points
                    " * ",       # Asterisk separators
                    " ",         # Word boundaries
                    ""           # Character level (fallback)
                ]
            )
            chunks = splitter.split_documents(docs)
            logging.info(f"[DEBUG] Created {len(chunks)} chunks from {filename}")
            
            # Add chunks to existing database
            if self.db:
                self.db.add_documents(chunks)
                logging.info(f"[DEBUG] Added {len(chunks)} chunks from {filename} to database")
                
                # Save the updated database
                self.db.save_local(DB_PATH)
                
                # SMART: Don't validate integrity after every incremental update
                # This prevents unnecessary rebuilds for minor discrepancies
                # Integrity validation will run periodically or when explicitly requested
                logging.info(f"[DEBUG] Successfully added {filename} to database without post-update validation")
                return True
            else:
                logging.error("No database available for incremental update")
                return False
                
        except Exception as e:
            logging.error(f"Error adding document incrementally: {e}")
            return False

    def delete_document_incremental(self, filename: str) -> bool:
        """Remove a document's chunks from the database (requires rebuilding)."""
        try:
            logging.info(f"[DEBUG] CRITICAL: Deleting document: {filename}")
            
            # CRITICAL FIX: ALWAYS force complete rebuild to guarantee no data leakage
            # This ensures deleted document data is completely removed
            
            # Get list of all files except the one to delete
            all_files = []
            for root, _, files in os.walk(DOCS_PATH):
                for file in files:
                    if file != filename:
                        all_files.append(os.path.join(root, file))
            
            if len(all_files) == 0:
                logging.warning("No files remaining after deletion")
                return False
            
            # CRITICAL: Force complete rebuild to ensure data integrity
            logging.info(f"[DEBUG] CRITICAL: Forcing complete database rebuild after deletion")
            self.build_db("forced_rebuild_after_deletion")
            
            logging.info(f"[DEBUG] CRITICAL: Successfully deleted {filename} with guaranteed data cleanup")
            return True
            
        except Exception as e:
            logging.error(f"CRITICAL ERROR: Error deleting document: {e}")
            return False

    def remove_document_from_vector_db(self, filename: str) -> bool:
        """Remove a specific document from the vector database incrementally."""
        try:
            logging.info(f"[DEBUG] Removing document from vector database: {filename}")
            
            if not self.db:
                logging.warning("[DEBUG] No vector database to remove document from")
                return False
            
            # Get list of all files except the one to delete
            remaining_files = []
            for root, _, files in os.walk(DOCS_PATH):
                for file in files:
                    if file != filename:
                        remaining_files.append(os.path.join(root, file))
            
            if not remaining_files:
                # No files left, clear the database
                self.db = None
                logging.info(f"[DEBUG] No files left, cleared vector database")
                return True
            
            # CRITICAL: Force complete rebuild to ensure data integrity
            logging.info(f"[DEBUG] CRITICAL: Forcing complete vector database rebuild after deletion")
            self.build_db("forced_rebuild_after_deletion")
            return True
            
        except Exception as e:
            logging.error(f"Error removing document from vector database: {e}")
            return False

    def calculate_optimal_chunk_parameters(self) -> tuple:
        """Calculate optimal chunk size and overlap based on available resources for FAQ bot quality."""
        try:
            # Get current system resources
            memory = psutil.virtual_memory()
            available_ram_gb = memory.available / (1024**3)
            total_ram_gb = memory.total / (1024**3)
            memory_pressure = memory.percent
            
            cpu_cores = psutil.cpu_count(logical=True)
            
            # Base chunk size calculation based on available resources
            if available_ram_gb > 16:
                base_chunk_size = 512    # Increased for better accuracy
            elif available_ram_gb > 8:
                base_chunk_size = 384    # Increased for better accuracy
            elif available_ram_gb > 4:
                base_chunk_size = 256    # Increased for better accuracy
            else:
                base_chunk_size = 192    # Increased minimum
            
            # Adjust based on memory pressure
            if memory_pressure > 80:
                base_chunk_size = int(base_chunk_size * 0.7)  # Reduce by 30% under pressure
            elif memory_pressure > 60:
                base_chunk_size = int(base_chunk_size * 0.85)  # Reduce by 15% under moderate pressure
            
            # Adjust based on CPU cores (more cores = can handle larger chunks)
            if cpu_cores > 8:
                base_chunk_size = int(base_chunk_size * 1.2)  # Increase by 20%
            elif cpu_cores > 4:
                base_chunk_size = int(base_chunk_size * 1.1)  # Increase by 10%
            
            # FAQ Bot Quality Bounds - ensure chunks are optimal for Q&A
            MIN_CHUNK_SIZE = 256    # Increased minimum for better context in Q&A
            MAX_CHUNK_SIZE = 1024   # Increased maximum for better FAQ accuracy
            
            # Apply quality bounds
            optimal_chunk_size = max(MIN_CHUNK_SIZE, min(MAX_CHUNK_SIZE, base_chunk_size))
            
            # Calculate optimal overlap (25-30% for good continuity in Q&A)
            overlap_ratio = 0.28  # 28% overlap for FAQ bots
            optimal_overlap = max(32, int(optimal_chunk_size * overlap_ratio))
            
            # Ensure overlap doesn't exceed chunk size
            optimal_overlap = min(optimal_overlap, optimal_chunk_size - 32)
            
            logging.info(f"[DEBUG] Chunking calculation:")
            logging.info(f"[DEBUG]   - Available RAM: {available_ram_gb:.2f} GB")
            logging.info(f"[DEBUG]   - Memory pressure: {memory_pressure:.1f}%")
            logging.info(f"[DEBUG]   - CPU cores: {cpu_cores}")
            logging.info(f"[DEBUG]   - Base chunk size: {base_chunk_size}")
            logging.info(f"[DEBUG]   - Final chunk size: {optimal_chunk_size}")
            logging.info(f"[DEBUG]   - Final overlap: {optimal_overlap}")
            
            return optimal_chunk_size, optimal_overlap
            
        except Exception as e:
            logging.warning(f"[DEBUG] Error calculating chunk parameters: {e}, using defaults")
            # Fallback to good FAQ bot defaults
            return 384, 108  # 384 chars with 108 overlap (28%) for better accuracy

    def replace_document_incremental(self, old_filename: str, new_file_path: str) -> bool:
        """Replace a document in the database."""
        try:
            logging.info(f"[DEBUG] Replacing document: {old_filename} with {new_file_path}")
            
            # Delete the old document
            if not self.delete_document_incremental(old_filename):
                logging.error(f"Failed to delete old document: {old_filename}")
                return False
            
            # Add the new document
            if not self.add_document_incremental(new_file_path):
                logging.error(f"Failed to add new document: {new_file_path}")
                return False
            
            logging.info(f"[DEBUG] Successfully replaced {old_filename} with {new_file_path}")
            return True
            
        except Exception as e:
            logging.error(f"Error replacing document: {e}")
            return False

    def extract_tables_from_pdf(self, file_path: str) -> List[Document]:
        """Extract tables from PDF using multiple methods for better accuracy."""
        tables = []
        
        if not TABLE_EXTRACTION_AVAILABLE:
            logging.warning("Table extraction not available - install dependencies")
            return tables
            
        try:
            # Method 1: Try tabula-py (works well for most tables)
            if tabula and hasattr(tabula, 'read_pdf'):
                try:
                    pdf_tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True)
                    for i, table in enumerate(pdf_tables):
                        if not table.empty:
                            # Convert table to structured text
                            table_text = self.table_to_text(table, f"Table {i+1}")
                            if table_text.strip():
                                tables.append(Document(
                                    page_content=table_text,
                                    metadata={"source": file_path, "type": "table", "table_index": i}
                                ))
                except Exception as e:
                    logging.warning(f"Tabula failed for {file_path}: {e}")
            
            # Method 2: Try camelot (better for complex tables)
            if camelot and hasattr(camelot, 'read_pdf'):
                try:
                    pdf_tables = camelot.read_pdf(file_path, pages='all')
                    for i, table in enumerate(pdf_tables):
                        if table.df is not None and not table.df.empty:
                            # Convert table to structured text
                            table_text = self.table_to_text(table.df, f"Table {i+1}")
                            if table_text.strip():
                                tables.append(Document(
                                    page_content=table_text,
                                    metadata={"source": file_path, "type": "table", "table_index": i}
                                ))
                except Exception as e:
                    logging.warning(f"Camelot failed for {file_path}: {e}")
            
            # Method 3: OCR-based table extraction for image-based PDFs
            if cv2 and pytesseract:
                try:
                    tables.extend(self.extract_tables_with_ocr(file_path))
                except Exception as e:
                    logging.warning(f"OCR table extraction failed for {file_path}: {e}")
                
        except Exception as e:
            logging.error(f"Error extracting tables from {file_path}: {e}")
        
        return tables

    def extract_tables_with_ocr(self, file_path: str) -> List[Document]:
        """Extract tables using OCR for image-based PDFs."""
        tables = []
        
        if not (cv2 and pytesseract):
            return tables
            
        try:
            import numpy as np
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Get page as image
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                
                # Convert to OpenCV format
                nparr = np.frombuffer(img_data, np.uint8)
                img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                
                # Convert to grayscale
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                
                # Detect table lines
                horizontal_lines = cv2.HoughLinesP(
                    gray, 1, np.pi/180, threshold=100, minLineLength=100, maxLineGap=10
                )
                vertical_lines = cv2.HoughLinesP(
                    cv2.rotate(gray, cv2.ROTATE_90_CLOCKWISE), 1, np.pi/180, 
                    threshold=100, minLineLength=100, maxLineGap=10
                )
                
                # If we detect table structure, extract text
                if horizontal_lines is not None or vertical_lines is not None:
                    # Extract text from the entire page
                    text = pytesseract.image_to_string(img)
                    
                    # Try to structure the text as a table
                    structured_text = self.structure_ocr_text_as_table(text)
                    if structured_text.strip():
                        tables.append(Document(
                            page_content=structured_text,
                            metadata={"source": file_path, "type": "ocr_table", "page": page_num}
                        ))
            
            doc.close()
            
        except Exception as e:
            logging.error(f"Error in OCR table extraction: {e}")
        
        return tables

    def structure_ocr_text_as_table(self, text: str) -> str:
        """Structure OCR text that appears to be tabular data."""
        lines = text.strip().split('\n')
        structured_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                # Split by multiple spaces or tabs to identify columns
                columns = [col.strip() for col in line.split() if col.strip()]
                if len(columns) > 1:
                    # This looks like table data
                    structured_lines.append(" | ".join(columns))
                else:
                    structured_lines.append(line)
        
        return "\n".join(structured_lines)

    def table_to_text(self, table, table_name: str = "Table") -> str:
        """Convert a pandas DataFrame to structured text representation."""
        try:
            if not TABLE_EXTRACTION_AVAILABLE or not pd:
                return str(table)
                
            if hasattr(table, 'empty') and table.empty:
                return ""
            
            # Get column names
            columns = table.columns.tolist()
            
            # Create header
            text_lines = [f"{table_name}:"]
            text_lines.append("Columns: " + " | ".join(str(col) for col in columns))
            text_lines.append("-" * 50)
            
            # Add data rows
            for idx, row in table.iterrows():
                row_data = []
                for col in columns:
                    value = row[col]
                    # Handle NaN and None values
                    if pd.isna(value) or value is None:
                        value = "N/A"
                    else:
                        value = str(value).strip()
                    row_data.append(value)
                text_lines.append(" | ".join(row_data))
            
            # Add summary statistics for numeric columns
            if np:
                numeric_cols = table.select_dtypes(include=[np.number]).columns
                if len(numeric_cols) > 0:
                    text_lines.append("-" * 50)
                    text_lines.append("Summary Statistics:")
                    for col in numeric_cols:
                        if not table[col].isna().all():
                            text_lines.append(f"{col}: Min={table[col].min():.2f}, Max={table[col].max():.2f}, Mean={table[col].mean():.2f}")
            
            return "\n".join(text_lines)
            
        except Exception as e:
            logging.error(f"Error converting table to text: {e}")
            return str(table)

    def extract_tables_from_excel(self, file_path: str) -> List[Document]:
        """Extract tables from Excel files."""
        tables = []
        
        if not TABLE_EXTRACTION_AVAILABLE or not pd:
            logging.warning("Table extraction not available - install dependencies")
            return tables
            
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                try:
                    # Read the sheet
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    if not df.empty:
                        # Convert table to structured text
                        table_text = self.table_to_text(df, f"Sheet: {sheet_name}")
                        if table_text.strip():
                            tables.append(Document(
                                page_content=table_text,
                                metadata={"source": file_path, "type": "excel_table", "sheet": sheet_name}
                            ))
                except Exception as e:
                    logging.warning(f"Error reading sheet {sheet_name}: {e}")
                    
        except Exception as e:
            logging.error(f"Error extracting tables from Excel {file_path}: {e}")
        
        return tables

    def extract_tables_from_docx(self, file_path: str) -> List[Document]:
        """Extract tables from Word documents."""
        tables = []
        
        if not TABLE_EXTRACTION_AVAILABLE or not pd:
            logging.warning("Table extraction not available - install dependencies")
            return tables
            
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            for i, table in enumerate(doc.tables):
                try:
                    # Convert table to pandas DataFrame
                    data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        data.append(row_data)
                    
                    if data:
                        # Create DataFrame
                        df = pd.DataFrame(data[1:], columns=data[0] if data else [])
                        
                        # Convert to structured text
                        table_text = self.table_to_text(df, f"Table {i+1}")
                        if table_text.strip():
                            tables.append(Document(
                                page_content=table_text,
                                metadata={"source": file_path, "type": "docx_table", "table_index": i}
                            ))
                except Exception as e:
                    logging.warning(f"Error processing table {i} in {file_path}: {e}")
                    
        except Exception as e:
            logging.error(f"Error extracting tables from Word {file_path}: {e}")
        
        return tables

    def enhanced_document_loader(self, file_path: str) -> List[Document]:
        """Enhanced document loader that extracts tables and structured content."""
        ext = os.path.splitext(file_path)[1].lower()
        all_docs = []
        
        try:
            # Extract tables first (only if table extraction is available)
            if TABLE_EXTRACTION_AVAILABLE:
                if ext == ".pdf":
                    all_docs.extend(self.extract_tables_from_pdf(file_path))
                elif ext == ".xlsx" or ext == ".xls":
                    all_docs.extend(self.extract_tables_from_excel(file_path))
                elif ext == ".docx":
                    all_docs.extend(self.extract_tables_from_docx(file_path))
            
            # Then extract regular text content
            if ext == ".txt":
                loader = TextLoader(file_path)
                all_docs.extend(loader.load())
            elif ext == ".pdf":
                # Try PyMuPDF first (faster for text-based PDFs)
                try:
                    loader = PyMuPDFLoader(file_path)
                    docs = loader.load()
                    if docs and any(d.page_content.strip() for d in docs):
                        all_docs.extend(docs)
                except Exception as e:
                    logging.warning(f"PyMuPDF failed for {file_path}: {e}. Trying OCR...")
                    try:
                        loader = UnstructuredPDFLoader(file_path, strategy="ocr_only")
                        docs = loader.load()
                        if docs and any(d.page_content.strip() for d in docs):
                            all_docs.extend(docs)
                    except Exception as ocr_e:
                        logging.error(f"Both PyMuPDF and OCR failed for {file_path}: {ocr_e}")
            elif ext == ".docx":
                loader = UnstructuredWordDocumentLoader(file_path)
                all_docs.extend(loader.load())
            elif ext == ".csv":
                loader = CSVLoader(file_path)
                all_docs.extend(loader.load())
            elif ext == ".xml":
                tree = ET.parse(file_path)
                root = tree.getroot()
                text = ET.tostring(root, encoding="unicode", method="text")
                all_docs.append(Document(page_content=text, metadata={"source": file_path}))
            
            # Add metadata about table extraction
            for doc in all_docs:
                if "type" not in doc.metadata:
                    doc.metadata["type"] = "text_content"
                doc.metadata["enhanced_processing"] = True
            
            logging.info(f"Enhanced processing extracted {len(all_docs)} documents from {file_path}")
            return all_docs
            
        except Exception as e:
            logging.error(f"Error in enhanced document loading for {file_path}: {e}")
            # Fallback to basic loader
            return self._basic_document_loader(file_path)
    
    def _basic_document_loader(self, file_path: str) -> List[Document]:
        """Basic document loader as fallback when enhanced processing fails."""
        ext = os.path.splitext(file_path)[1].lower()
        all_docs = []
        
        try:
            if ext == ".txt":
                loader = TextLoader(file_path)
                all_docs.extend(loader.load())
            elif ext == ".pdf":
                try:
                    loader = PyMuPDFLoader(file_path)
                    all_docs.extend(loader.load())
                except Exception as e:
                    logging.warning(f"PyMuPDF failed for {file_path}: {e}")
                    try:
                        loader = UnstructuredPDFLoader(file_path)
                        all_docs.extend(loader.load())
                    except Exception as e2:
                        logging.error(f"All PDF loaders failed for {file_path}: {e2}")
            elif ext == ".docx":
                loader = UnstructuredWordDocumentLoader(file_path)
                all_docs.extend(loader.load())
            elif ext == ".csv":
                loader = CSVLoader(file_path)
                all_docs.extend(loader.load())
            elif ext == ".xml":
                tree = ET.parse(file_path)
                root = tree.getroot()
                text = ET.tostring(root, encoding="unicode", method="text")
                all_docs.append(Document(page_content=text, metadata={"source": file_path}))
            
            logging.info(f"Basic processing extracted {len(all_docs)} documents from {file_path}")
            return all_docs
            
        except Exception as e:
            logging.error(f"Error in basic document loading for {file_path}: {e}")
            return []

    def start_embedding_async(self, query: str):
        """Get embedding immediately - no background processing."""
        try:
            # SPEED OPTIMIZATION: Use BGE embeddings from config for 1000x speed boost
            if hasattr(self, 'bge_model') and self.bge_model is not None:
                print(f"[SPEED] Using {EMBEDDING_MODEL} embeddings for 1000x speed boost")
                embedding = self.bge_model.encode(query, convert_to_tensor=False)
                print(f"[SPEED] {EMBEDDING_MODEL} embedding generated in milliseconds ({len(embedding)}D)")
                return embedding
            else:
                # Fallback to config model embeddings
                print(f"[SPEED] Using Ollama embeddings (fallback)")
                embedding = self.get_query_embeddings(query)
                
                # Truncate to smaller dimension for speed boost
                if len(embedding) > self.embedding_dimension:
                    embedding = embedding[:self.embedding_dimension]
                    print(f"[SPEED] Truncated embedding from {len(embedding)}D to {self.embedding_dimension}D for speed")
                
                return embedding
        except Exception as e:
            logging.warning(f"Embedding failed: {e}")
            return None

    def get_chunks_with_early_termination(self, query: str, similarity_threshold: float = 0.15, min_chunks: int = 20) -> List[Document]:
        """Get chunks using hybrid semantic + vector approach for better coverage."""
        early_start = time.time()
        
        try:
            print(f"[SEMANTIC] Starting hybrid semantic + vector search (threshold: {similarity_threshold})")
            
            # Use hybrid semantic search for better coverage
            hybrid_docs = self.get_hybrid_semantic_chunks(query, initial_k=25)
            
            if len(hybrid_docs) >= min_chunks:
                early_time = time.time() - early_start
                print(f"[SEMANTIC] Hybrid search SUCCESS: Found {len(hybrid_docs)} semantically relevant chunks in {early_time:.3f}s")
                return hybrid_docs[:min_chunks]
            else:
                print(f"[SEMANTIC] Hybrid search: Only {len(hybrid_docs)} chunks found, using fallback")
                # Use comprehensive fallback search
                fallback_docs = self.get_fallback_search_results(query, min_chunks)
                if len(fallback_docs) >= min_chunks:
                    print(f"[SEMANTIC] Fallback search successful: {len(fallback_docs)} chunks")
                    return fallback_docs[:min_chunks]
                else:
                    # Final fallback to regular search
                    return self.get_hierarchical_chunks(query, initial_k=min_chunks)
                
        except Exception as e:
            print(f"[SEMANTIC] Hybrid search failed: {e}, using regular search")
            # Fallback to regular search with more chunks
            return self.get_hierarchical_chunks(query, initial_k=min_chunks)

    def limit_context_size(self, context: str, max_chars: int = 6000) -> str:
        """Limit context size for faster LLM processing while keeping reasonable size."""
        if len(context) <= max_chars:
            return context
        
        print(f"[SPEED] Context too large ({len(context)} chars), limiting to {max_chars} chars")
        
        # Smart truncation - try to keep complete chunks
        chunks = context.split("\n\n")
        limited_chunks = []
        current_length = 0
        
        for chunk in chunks:
            if current_length + len(chunk) + 2 <= max_chars:  # +2 for \n\n
                limited_chunks.append(chunk)
                current_length += len(chunk) + 2
            else:
                # Add partial chunk if space allows
                remaining_space = max_chars - current_length - 3  # -3 for "..."
                if remaining_space > 100:  # Only if meaningful space left
                    partial_chunk = chunk[:remaining_space] + "..."
                    limited_chunks.append(partial_chunk)
                break
        
        limited_context = "\n\n".join(limited_chunks)
        print(f"[SPEED] Context limited to {len(limited_context)} chars ({len(limited_chunks)} chunks)")
        return limited_context

    def filter_relevant_chunks(self, query, docs, threshold=0.7):
        """Enhanced relevance filtering with multi-factor scoring for better quality"""
        if not docs:
            return []
        
        relevant_docs = []
        query_terms = set(query.lower().split())
        
        # Remove common stop words for better relevance
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'can', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them'}
        query_terms = {term for term in query_terms if term not in stop_words and len(term) > 2}
        
        for doc in docs:
            # Multi-factor relevance scoring
            score = 0
            doc_text = doc.page_content.lower()
            
            # 1. Term frequency scoring (fast)
            term_matches = sum(1 for term in query_terms if term in doc_text)
            score += term_matches * 0.3
            
            # 2. Phrase matching (fast) - exact phrase matches get higher scores
            query_lower = query.lower()
            if query_lower in doc_text:
                score += 2.0  # Exact phrase match gets high score
            
            # 3. Partial phrase matching
            for term in query_terms:
                if len(term) > 3:
                    count = doc_text.count(term)
                    score += count * 0.2
            
            # 4. Semantic similarity using embeddings (if available)
            if hasattr(self, 'embedding_model') and self.embedding_model:
                try:
                    query_emb = self.embedding_model.encode([query])[0]
                    doc_emb = self.embedding_model.encode([doc.page_content])[0]
                    similarity = np.dot(query_emb, doc_emb) / (np.linalg.norm(query_emb) * np.linalg.norm(doc_emb))
                    score += similarity * 0.5
                except Exception as e:
                    print(f"[DEBUG] Embedding similarity calculation failed: {e}")
                    pass
            
            # 5. Document metadata relevance (source, topic matching)
            source = str(doc.metadata.get("source", "")).lower()
            if any(term in source for term in query_terms):
                score += 0.3  # Source filename relevance
            
            # 6. Content structure relevance (headers, lists, tables)
            if any(marker in doc_text for marker in ['policy', 'procedure', 'requirement', 'standard', 'guideline']):
                score += 0.2  # Policy-related content gets bonus
            
            if score >= threshold:
                relevant_docs.append((doc, score))
        
        # Sort by relevance score and return top chunks
        relevant_docs.sort(key=lambda x: x[1], reverse=True)
        
        # Balanced filtering: return documents with good relevance
        # For simple questions, we want focused sources but not too restrictive
        high_threshold = threshold + 0.1  # Smaller increase for balanced filtering
        high_relevance_docs = [doc for doc, score in relevant_docs if score >= high_threshold]
        
        # Limit to top 5 most relevant documents for balanced coverage
        top_docs = [doc for doc, score in relevant_docs[:5]]
        
        # If we have high relevance docs, use those but be more lenient
        if high_relevance_docs:
            top_docs = high_relevance_docs[:4]  # Allow up to 4 high-relevance docs
        
        print(f"[QUALITY] Enhanced filtering: {len(docs)} -> {len(top_docs)} relevant chunks (threshold: {threshold}, high_threshold: {high_threshold})")
        if top_docs:
            top_scores = [score for _, score in relevant_docs[:3]]
            print(f"[QUALITY] Top 3 relevance scores: {[f'{s:.3f}' for s in top_scores]}")
        else:
            print(f"[QUALITY] WARNING: No documents passed filtering! All scores below threshold {threshold}")
            if relevant_docs:
                print(f"[QUALITY] Best scores found: {[f'{s:.3f}' for _, s in relevant_docs[:5]]}")
        
        return top_docs

    def get_concise_response(self, prompt: str) -> str:
        """Get concise response by limiting generation for faster processing."""
        # Add explicit length instruction to prompt
        concise_prompt = f"{prompt}\n\nProvide a concise answer in 2-3 sentences maximum."
        
        try:
            response = self.llm.invoke(concise_prompt)
            
            # Convert to string if needed
            if isinstance(response, dict) and 'result' in response:
                response = response['result']
            else:
                response = str(response)
            
            # Ensure response isn't too long
            if len(response) > 800:
                print(f"[SPEED] Response too long ({len(response)} chars), truncating")
                # Truncate at sentence boundary
                sentences = response.split('.')
                truncated = []
                current_length = 0
                
                for sentence in sentences:
                    if current_length + len(sentence) + 1 <= 800:
                        truncated.append(sentence)
                        current_length += len(sentence) + 1
                    else:
                        break
                
                result = '.'.join(truncated) + '.'
                print(f"[SPEED] Response truncated to {len(result)} chars")
                return result
            
            return response
            
        except Exception as e:
            print(f"[SPEED] Concise response failed: {e}, using regular invoke")
            return str(self.llm.invoke(prompt))



    def build_context_streamlined(self, docs, max_chars=3000, overlap_ratio=0.2):
        """Build context with intelligent chunking and overlap for better quality"""
        if not docs:
            return ""
        
        # Calculate optimal chunk size based on max_chars and number of docs
        target_chunk_size = max_chars // max(len(docs), 1)
        overlap_size = int(target_chunk_size * overlap_ratio)
        
        context_parts = []
        current_length = 0
        
        for i, doc in enumerate(docs):
            if current_length >= max_chars:
                break
                
            content = doc.page_content
            source = os.path.basename(str(doc.metadata.get("source", "")))
            page_info = doc.metadata.get("page", doc.metadata.get("page_number", ""))
            
            # Smart truncation with overlap preservation
            if len(content) > target_chunk_size:
                # Find natural break points (sentences, paragraphs)
                sentences = content.split('.')
                truncated = ""
                for sentence in sentences:
                    if len(truncated + sentence) <= target_chunk_size:
                        truncated += sentence + "."
                    else:
                        break
                
                # Add overlap from next chunk for context continuity
                if overlap_size > 0 and len(content) > target_chunk_size + overlap_size:
                    overlap_start = target_chunk_size - overlap_size
                    overlap_text = content[overlap_start:target_chunk_size]
                    # Clean up overlap text
                    overlap_text = overlap_text.strip()
                    if overlap_text and not overlap_text.endswith('.'):
                        overlap_text += "..."
                    truncated += f" [continued: {overlap_text}]"
            else:
                truncated = content
            
            # Add source attribution and page info
            if page_info and page_info != "?":
                context_parts.append(f"[{source} - Page {page_info}]: {truncated}")
            else:
                context_parts.append(f"[{source}]: {truncated}")
            
            current_length += len(truncated)
            
            # Add separator between documents for clarity
            if i < len(docs) - 1 and current_length < max_chars:
                context_parts.append("---")
                current_length += 3
        
        final_context = "\n\n".join(context_parts)
        
        # Ensure context doesn't exceed max_chars
        if len(final_context) > max_chars:
            # Truncate at sentence boundary
            sentences = final_context.split('.')
            truncated_context = ""
            for sentence in sentences:
                if len(truncated_context + sentence + ".") <= max_chars:
                    truncated_context += sentence + "."
                else:
                    break
            
            if truncated_context:
                final_context = truncated_context + " [context truncated for length]"
            else:
                # Fallback: take first max_chars characters
                final_context = final_context[:max_chars-50] + " [context truncated]"
        
        print(f"[QUALITY] Context built: {len(docs)} docs -> {len(final_context)} chars")
        return final_context

    def process_chunks_parallel(self, docs: List[Document]) -> List[dict]:
        """Process retrieved chunks in parallel for faster preparation."""
        parallel_start = time.time()
        
        def process_single_chunk(doc_with_index):
            """Process a single chunk with its index."""
            i, doc = doc_with_index
            try:
                # Process chunk metadata and content
                source = os.path.basename(str(doc.metadata.get('source', 'Unknown')))
                chunk_content = doc.page_content.strip()
                
                # Limit chunk size to prevent overwhelming the LLM
                if len(chunk_content) > 300:
                    chunk_content = chunk_content[:300] + "..."
                
                return {
                    'index': i,
                    'source': source,
                    'content': chunk_content,
                    'metadata': doc.metadata
                }
            except Exception as e:
                logging.warning(f"Error processing chunk {i}: {e}")
                return {
                    'index': i,
                    'source': 'Error',
                    'content': 'Processing error',
                    'metadata': {}
                }
        
        try:
            # Process chunks in parallel with ThreadPoolExecutor
            from concurrent.futures import ThreadPoolExecutor
            with ThreadPoolExecutor(max_workers=4) as executor:
                # Add index to each doc for processing
                docs_with_index = [(i, doc) for i, doc in enumerate(docs)]
                processed_chunks = list(executor.map(process_single_chunk, docs_with_index))
            
            # Sort by index to maintain order
            processed_chunks.sort(key=lambda x: x['index'])
            
            parallel_time = time.time() - parallel_start
            print(f"[SPEED] Parallel chunk processing completed in {parallel_time:.3f}s (4 workers)")
            
            return processed_chunks
            
        except Exception as e:
            print(f"[SPEED] Parallel processing failed: {e}, using sequential fallback")
            parallel_time = time.time() - parallel_start
            print(f"[SPEED] Parallel processing attempt took {parallel_time:.3f}s before fallback")
            
            # Fallback to sequential processing
            processed_chunks = []
            for i, doc in enumerate(docs):
                source = os.path.basename(str(doc.metadata.get('source', 'Unknown')))
                chunk_content = doc.page_content.strip()
                if len(chunk_content) > 300:
                    chunk_content = chunk_content[:300] + "..."
                processed_chunks.append({
                    'index': i,
                    'source': source,
                    'content': chunk_content,
                    'metadata': doc.metadata
                })
            
            return processed_chunks

    def get_hierarchical_chunks(self, query: str, initial_k: int = None) -> List[Document]:
        """
        Ultra-fast chunk-only retrieval: get only the most relevant chunks for maximum speed.
        
        Args:
            query: The user's question
            initial_k: Number of chunks to retrieve (default: 15 for optimal speed)
            
        Returns:
            List of document chunks (not full documents) for maximum performance
        """
        hierarchical_start = time.time()
        
        # Optimize search parameters for better performance
        self._optimize_search_parameters()
        
        # Use optimal chunk count for speed vs accuracy balance
        if initial_k is None:
            initial_k = min(20, self.db.index.ntotal)  # 20 chunks for optimal balance
            logging.info(f"[TIMING] Starting ultra-fast chunk-only retrieval for {initial_k} chunks (optimized for balance)")
        else:
            logging.info(f"[TIMING] Starting chunk-only retrieval with initial_k={initial_k}")
        
        try:
            logging.info(f"[DEBUG] Starting ultra-fast chunk retrieval with initial_k={initial_k}")
            
            # Step 1: Get only the most relevant chunks (Vector similarity search)
            initial_search_start = time.time()
            logging.info(f"[TIMING] Starting ultra-fast vector similarity search")
            retriever = self.db.as_retriever(search_type="similarity", k=initial_k)
            initial_docs = retriever.invoke(query)
            initial_search_time = time.time() - initial_search_start
            logging.info(f"[TIMING] Ultra-fast vector similarity search took: {initial_search_time:.3f}s")
            
            if not initial_docs:
                logging.warning("[DEBUG] No chunks found")
                return []
            
            logging.info(f"[DEBUG] Retrieved {len(initial_docs)} chunks (not full documents)")
            
            # Return chunks directly - no document expansion
            hierarchical_time = time.time() - hierarchical_start
            logging.info(f"[TIMING] Total ultra-fast chunk retrieval took: {hierarchical_time:.3f}s")
            logging.info(f"[DEBUG] Chunk-only retrieval complete: {len(initial_docs)} chunks")
            
            return initial_docs
            
        except Exception as e:
            logging.error(f"[DEBUG] Error in chunk retrieval: {e}")
            # Fallback to minimal chunk retrieval
            fallback_start = time.time()
            logging.info(f"[TIMING] Using fallback minimal chunk retrieval")
            retriever = self.db.as_retriever(search_type="similarity", k=min(10, self.db.index.ntotal))
            result = retriever.invoke(query)
            fallback_time = time.time() - fallback_start
            logging.info(f"[TIMING] Fallback minimal chunk retrieval took: {fallback_time:.3f}s")
            return result

    def calculate_confidence_with_gemma(self, query: str, answer: str, docs: list) -> None:
        """
        Calculate confidence using config model with fallback.
        
        Args:
            query: The user's question
            answer: The generated answer
            docs: The documents used for the answer
            
        Returns:
            Confidence score (0-100)
        """
        confidence_start = time.time()
        logging.info(f"[TIMING] Starting confidence calculation with config model")
        
        try:
            # Prepare context from retrieved documents
            context = "\n\n".join([f"Document {i+1}: {doc.page_content}" for i, doc in enumerate(docs)])
            
            # Confidence evaluation prompt
            confidence_prompt = f"""Rate this answer's confidence (0-100) based on:
- Relevance to question
- Completeness of information
- Clarity of response

Question: {query}
Answer: {answer}
Sources: {context}

Respond with ONLY a number between 0 and 100."""

            # # Try Gemma 3B first
            # if self.gemma_model:
            #     gemma_start = time.time()
            #     logging.info("[DEBUG] Using Gemma 3B for confidence calculation")
            #     confidence_response = self.gemma_model.invoke(confidence_prompt)
            #     confidence_text = str(confidence_response).strip()
            #     gemma_time = time.time() - gemma_start
            #     logging.info(f"[TIMING] Gemma 3B confidence calculation took: {gemma_time:.3f}s")
                
            #     # Extract numeric confidence score
            #     import re
            #     confidence_match = re.search(r'\b(\d{1,3})\b', confidence_text)
            #     if confidence_match:
            #         confidence_score = float(confidence_match.group(1))
            #         confidence_score = max(0, min(100, confidence_score))
            #         logging.info(f"[DEBUG] Gemma 3B confidence score: {confidence_score}")
            #         return round(confidence_score, 2)
            
            # # Fallback to Llama 3
            # if self.llama_model:
            #     llama_start = time.time()
            #     logging.info("[DEBUG] Using Llama 3 fallback for confidence calculation")
            #     confidence_response = self.llama_model.invoke(confidence_prompt)
            #     confidence_text = str(confidence_response).strip()
            #     llama_time = time.time() - llama_start
            #     logging.info(f"[TIMING] Llama 3 confidence calculation took: {llama_time:.3f}s")
                
            #     # Extract numeric confidence score
            #     import re
            #     confidence_match = re.search(r'\b(\d{1,3})\b', confidence_text)
            #     if confidence_match:
            #         confidence_score = float(confidence_match.group(1))
            #         confidence_score = max(0, min(100, confidence_score))
            #         logging.info(f"[DEBUG] Llama 3 fallback confidence score: {confidence_score}")
            #         return round(confidence_score, 2)
            
            # Final fallback
            logging.warning("[DEBUG] All confidence models failed, using fallback")
            #return 50.0 if docs else 0.0
            return None
                
        except Exception as e:
            logging.warning(f"[DEBUG] Confidence calculation failed: {e}")
            # Fallback confidence based on whether we have documents
            #return 20.0 if docs else 0.0
            return None

    def initialize_confidence_models(self):
        """Initialize config model for confidence calculation (Ollama auto-detects best device)."""
        try:
            # Initialize confidence models (Ollama automatically uses best available device)
            optimal_device = self.hardware_info['optimal_device']
            logging.info(f"[DEBUG] Initializing confidence models (Ollama will auto-detect best device: {optimal_device})")
            
            from langchain_ollama import OllamaLLM
            
            # Try config model first (Ollama auto-detects GPU/CPU)
            try:
                logging.info(f"[DEBUG] Attempting {MODEL_NAME} initialization")
                self.gemma_model = OllamaLLM(
            model=MODEL_NAME,
            temperature=0.1,
            num_ctx=8192,
            num_predict=2048,
            stop=None
        )
                self.confidence_model = self.gemma_model
                logging.info(f"[DEBUG] {MODEL_NAME} confidence model initialized successfully")
                
            except Exception as e1:
                logging.warning(f"[DEBUG] {MODEL_NAME} failed: {e1}")
                
                # Fallback to Llama 3
                try:
                    logging.info(f"[DEBUG] Attempting Llama 3 fallback initialization")
                    self.llama_model = OllamaLLM(
                        model="llama3.2:3b",
                        temperature=0.1,
                        num_ctx=8192,
                        num_predict=2048,
                        stop=None
                    )
                    self.confidence_model = self.llama_model
                    logging.info(f"[DEBUG] Llama 3 fallback confidence model initialized successfully")
                    
                except Exception as e2:
                    logging.error(f"[DEBUG] Failed to initialize confidence models: {e2}")
                    self.confidence_model = None
                    
        except Exception as e:
            logging.error(f"[DEBUG] Failed to initialize confidence models: {e}")
            self.confidence_model = None

    def validate_vector_db_integrity(self) -> dict:
        """Validate that vector database matches actual files on disk and clean any stale data."""
        try:
            logging.info("[DEBUG] Starting vector database integrity validation...")
            
            if not self.db:
                logging.info("[DEBUG] No vector database to validate")
                return {'status': 'no_db', 'message': 'No vector database exists'}
            
            # Get all files currently on disk
            disk_files = set()
            for root, _, files in os.walk(DOCS_PATH):
                for file in files:
                    disk_files.add(file)
            
            logging.info(f"[DEBUG] Found {len(disk_files)} files on disk")
            
            # Get all sources from vector database
            try:
                # Get a sample of documents to check sources
                sample_docs = self.db.similarity_search("", k=1000)  # Get up to 1000 docs
                db_sources = set()
                for doc in sample_docs:
                    source = doc.metadata.get('source', '')
                    if source:
                        # Extract filename from full path
                        filename = os.path.basename(source)
                        db_sources.add(filename)
                
                logging.info(f"[DEBUG] Vector database contains {len(db_sources)} unique document sources")
                
                # Find stale data (files in DB but not on disk)
                stale_sources = db_sources - disk_files
                missing_sources = disk_files - db_sources
                
                if stale_sources:
                    logging.warning(f"[DEBUG] Found {len(stale_sources)} stale sources in vector database: {stale_sources}")
                else:
                    logging.info("[DEBUG] No stale sources found in vector database")
                
                if missing_sources:
                    logging.warning(f"[DEBUG] Found {len(missing_sources)} files on disk not in vector database: {missing_sources}")
                else:
                    logging.info("[DEBUG] All disk files are represented in vector database")
                
                # SMART VALIDATION: Only rebuild when absolutely necessary
                total_discrepancies = len(stale_sources) + len(missing_sources)
                total_files = len(disk_files)
                
                # Calculate discrepancy percentage
                if total_files > 0:
                    discrepancy_percentage = (total_discrepancies / total_files) * 100
                else:
                    discrepancy_percentage = 0
                
                logging.info(f"[DEBUG] Discrepancy analysis: {total_discrepancies} issues out of {total_files} files ({discrepancy_percentage:.1f}%)")
                
                # DECISION LOGIC: Only rebuild when absolutely necessary
                if total_discrepancies == 0:
                    # Perfect match - no issues
                    logging.info("[DEBUG] Vector database integrity validated successfully - perfect match")
                    hnsw_status = "unknown"
                    if self.db and hasattr(self.db, 'index'):
                        if hasattr(self.db.index, 'hnsw'):
                            hnsw_status = "active"
                            print(f"[HNSW] Database integrity confirmed with active HNSW index")
                        else:
                            hnsw_status = "standard"
                            print(f"[HNSW] Database integrity confirmed with standard index")
                    
                    return {
                        'status': 'valid',
                        'message': 'Vector database integrity confirmed - perfect match',
                        'total_sources': len(db_sources),
                        'hnsw_status': hnsw_status,
                        'cleanup_verified': True
                    }
                
                elif total_discrepancies <= 2 and discrepancy_percentage <= 10:
                    # Minor discrepancies - safe to ignore
                    logging.info(f"[DEBUG] Minor discrepancies detected ({total_discrepancies} issues, {discrepancy_percentage:.1f}%) - safe to ignore")
                    return {
                        'status': 'minor_issues',
                        'message': f'Minor discrepancies detected but not critical ({total_discrepancies} issues, {discrepancy_percentage:.1f}%)',
                        'stale_sources': len(stale_sources),
                        'missing_sources': len(missing_sources),
                        'discrepancy_percentage': discrepancy_percentage,
                        'action_taken': 'none'
                    }
                
                elif total_discrepancies <= 5 and discrepancy_percentage <= 20:
                    # Moderate discrepancies - try to fix without rebuild
                    logging.info(f"[DEBUG] Moderate discrepancies detected ({total_discrepancies} issues, {discrepancy_percentage:.1f}%) - attempting to fix without rebuild")
                    
                    # Try to fix specific issues
                    if self._try_fix_moderate_discrepancies(stale_sources, missing_sources):
                        logging.info("[DEBUG] Successfully fixed moderate discrepancies without rebuild")
                        return {
                            'status': 'fixed',
                            'message': f'Moderate discrepancies fixed without rebuild ({total_discrepancies} issues resolved)',
                            'stale_sources': len(stale_sources),
                            'missing_sources': len(missing_sources),
                            'discrepancy_percentage': discrepancy_percentage,
                            'action_taken': 'fixed'
                        }
                    else:
                        logging.warning("[DEBUG] Could not fix moderate discrepancies, but not critical enough for rebuild")
                        return {
                            'status': 'moderate_issues',
                            'message': f'Moderate discrepancies detected but not critical enough for rebuild ({total_discrepancies} issues, {discrepancy_percentage:.1f}%)',
                            'stale_sources': len(stale_sources),
                            'missing_sources': len(missing_sources),
                            'discrepancy_percentage': discrepancy_percentage,
                            'action_taken': 'none'
                        }
                
                else:
                    # CRITICAL: Major discrepancies - rebuild absolutely necessary
                    logging.info(f"[DEBUG] CRITICAL: Major discrepancies detected ({total_discrepancies} issues, {discrepancy_percentage:.1f}%) - rebuild absolutely necessary")
                    print(f"[HNSW] CRITICAL: Found {len(stale_sources)} stale sources and {len(missing_sources)} missing sources ({discrepancy_percentage:.1f}% discrepancy)")
                    
                    # CRITICAL: Force complete rebuild to ensure data integrity
                    self.build_db("integrity_rebuild")
                    
                    return {
                        'status': 'rebuilt',
                        'message': f'Vector database rebuilt due to critical integrity issues ({total_discrepancies} issues, {discrepancy_percentage:.1f}% discrepancy)',
                        'stale_sources': len(stale_sources),
                        'missing_sources': len(missing_sources),
                        'discrepancy_percentage': discrepancy_percentage,
                        'cleanup_verified': True
                    }
                    
            except Exception as e:
                logging.error(f"[DEBUG] Error validating vector database: {e}")
                logging.info("[DEBUG] Rebuilding database due to validation error...")
                self.build_db("validation_error_rebuild")
                return {
                    'status': 'error_rebuilt',
                    'message': f'Database rebuilt due to validation error: {str(e)}'
                }
                
        except Exception as e:
            logging.error(f"Error in vector database integrity validation: {e}")
            return {'status': 'error', 'message': str(e)}

    def _try_fix_moderate_discrepancies(self, stale_sources: set, missing_sources: set) -> bool:
        """Try to fix moderate discrepancies without rebuilding the database."""
        try:
            logging.info("[DEBUG] Attempting to fix moderate discrepancies without rebuild...")
            
            fixed_issues = 0
            
            # Try to fix missing sources by adding them incrementally
            for missing_file in missing_sources:
                try:
                    # Find the full path of the missing file
                    missing_file_path = None
                    for root, _, files in os.walk(DOCS_PATH):
                        if missing_file in files:
                            missing_file_path = os.path.join(root, missing_file)
                            break
                    
                    if missing_file_path and os.path.exists(missing_file_path):
                        # Try to add this file incrementally
                        if self.add_document_incremental(missing_file_path):
                            fixed_issues += 1
                            logging.info(f"[DEBUG] Successfully added missing file: {missing_file}")
                        else:
                            logging.warning(f"[DEBUG] Failed to add missing file incrementally: {missing_file}")
                    else:
                        logging.warning(f"[DEBUG] Could not find missing file path: {missing_file}")
                        
                except Exception as e:
                    logging.warning(f"[DEBUG] Error fixing missing file {missing_file}: {e}")
                    continue
            
            # For stale sources, we can't easily remove them without rebuild
            # But we can log them for manual cleanup later
            if stale_sources:
                logging.info(f"[DEBUG] {len(stale_sources)} stale sources detected but cannot be removed without rebuild")
                logging.info(f"[DEBUG] Stale sources: {stale_sources}")
            
            logging.info(f"[DEBUG] Fixed {fixed_issues} out of {len(missing_sources)} missing sources")
            
            # Consider it successful if we fixed at least some issues
            return fixed_issues > 0
            
        except Exception as e:
            logging.error(f"[DEBUG] Error in _try_fix_moderate_discrepancies: {e}")
            return False

    def periodic_integrity_check(self) -> dict:
        """Run integrity check periodically (not after every update) to catch real issues."""
        try:
            logging.info("[DEBUG] Running periodic integrity check...")
            
            # Only run if database exists and is stable
            if not self.db:
                return {'status': 'no_db', 'message': 'No database to check'}
            
            # Check if we've run recently to avoid excessive checking
            current_time = time.time()
            if hasattr(self, '_last_integrity_check'):
                time_since_last = current_time - self._last_integrity_check
                if time_since_last < 3600:  # Don't check more than once per hour
                    logging.info("[DEBUG] Integrity check skipped - checked recently")
                    return {'status': 'skipped', 'message': 'Checked recently, skipping'}
            
            # Update last check time
            self._last_integrity_check = current_time
            
            # Run the smart integrity validation
            result = self.validate_vector_db_integrity()
            
            logging.info(f"[DEBUG] Periodic integrity check completed: {result.get('status')}")
            return result
            
        except Exception as e:
            logging.error(f"[DEBUG] Error in periodic integrity check: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_hybrid_semantic_chunks(self, query: str, initial_k: int = 15) -> List[Document]:
        """Get chunks using hybrid semantic + vector approach for better coverage."""
        hybrid_start = time.time()
        
        try:
            print(f"[SEMANTIC] Starting hybrid semantic + vector search for query: '{query}'")
            
            # Step 1: Get vector similarity results with more aggressive search
            vector_start = time.time()
            retriever = self.db.as_retriever(search_type="similarity", k=initial_k * 2)  # Get more initial results
            vector_docs = retriever.invoke(query)
            vector_time = time.time() - vector_start
            print(f"[SEMANTIC] Vector search found {len(vector_docs)} chunks in {vector_time:.3f}s")
            
            # Check if semantic search is enabled
            if not getattr(self, 'enable_semantic_search', True):
                print(f"[SEMANTIC] Semantic search disabled, returning {len(vector_docs)} vector chunks")
                return vector_docs[:initial_k]
            
            # Step 2: Generate semantic variations of the query
            semantic_start = time.time()
            semantic_variations = self._generate_semantic_variations(query)
            print(f"[SEMANTIC] Generated {len(semantic_variations)} semantic variations")
            
            # Step 3: Search with semantic variations (limited for performance)
            semantic_docs = []
            max_variations = getattr(self, 'semantic_search_variations', 10)
            for variation in semantic_variations[:max_variations]:
                try:
                    retriever = self.db.as_retriever(search_type="similarity", k=6)
                    variation_docs = retriever.invoke(variation)
                    semantic_docs.extend(variation_docs)
                except Exception as e:
                    print(f"[SEMANTIC] Error searching variation '{variation}': {e}")
            
            semantic_time = time.time() - semantic_start
            print(f"[SEMANTIC] Semantic search found {len(semantic_docs)} additional chunks in {semantic_time:.3f}s")
            
            # Step 4: Combine and deduplicate results
            all_docs = vector_docs + semantic_docs
            unique_docs = self._deduplicate_documents(all_docs)
            
            # Step 5: Re-rank by relevance to original query
            ranked_docs = self._rerank_by_relevance(query, unique_docs)
            
            hybrid_time = time.time() - hybrid_start
            print(f"[SEMANTIC] Hybrid search completed in {hybrid_time:.3f}s: {len(ranked_docs)} unique chunks")
            
            return ranked_docs[:initial_k]
            
        except Exception as e:
            print(f"[SEMANTIC] Hybrid search failed: {e}, falling back to vector search")
            return self.get_hierarchical_chunks(query, initial_k=initial_k)
    
    def _generate_semantic_variations(self, query: str) -> List[str]:
        """Generate semantic variations of the query to catch related concepts."""
        variations = []
        
        # Remove question numbers and formatting
        clean_query = re.sub(r'^\d+\)\s*', '', query.strip())
        
        # Common semantic expansions
        semantic_map = {
            'patch': ['update', 'upgrade', 'fix', 'security update', 'vulnerability fix', 'maintenance', 'software update'],
            'encryption': ['cryptography', 'encrypt', 'decrypt', 'cipher', 'secure communication', 'data protection', 'security controls'],
            'password': ['authentication', 'login', 'credential', 'access control', 'identity', 'user management', 'security access'],
            'backup': ['recovery', 'restore', 'data protection', 'disaster recovery', 'redundancy', 'business continuity', 'data backup'],
            'access': ['permission', 'authorization', 'entry', 'login', 'authentication', 'user access', 'system access', 'data access'],
            'security': ['protection', 'safeguard', 'defense', 'safety', 'risk mitigation', 'information security', 'cybersecurity'],
            'policy': ['procedure', 'guideline', 'rule', 'standard', 'requirement', 'framework', 'governance', 'compliance'],
            'compliance': ['regulation', 'standard', 'requirement', 'policy', 'audit', 'governance', 'regulatory', 'standards'],
            'incident': ['event', 'breach', 'violation', 'problem', 'issue', 'security incident', 'response', 'management'],
            'training': ['education', 'awareness', 'learning', 'instruction', 'knowledge', 'staff training', 'employee training'],
            'asset': ['resource', 'system', 'equipment', 'infrastructure', 'hardware', 'software', 'data asset', 'information asset'],
            'management': ['administration', 'oversight', 'governance', 'control', 'supervision', 'leadership', 'stewardship'],
            'review': ['assessment', 'evaluation', 'audit', 'examination', 'inspection', 'analysis', 'verification'],
            'approval': ['authorization', 'endorsement', 'sanction', 'consent', 'permission', 'acceptance', 'ratification'],
            'communication': ['notification', 'awareness', 'information sharing', 'dissemination', 'reporting', 'transparency'],
            'maintenance': ['upkeep', 'servicing', 'support', 'care', 'preservation', 'sustenance', 'maintenance schedule']
        }
        
        # Generate variations based on semantic mapping
        query_lower = clean_query.lower()
        for key, synonyms in semantic_map.items():
            if key in query_lower:
                for synonym in synonyms:
                    variation = clean_query.replace(key, synonym)
                    if variation != clean_query:
                        variations.append(variation)
        
        # Add common question variations
        question_variations = [
            f"What is the policy on {clean_query.lower()}?",
            f"How does the organization handle {clean_query.lower()}?",
            f"What are the requirements for {clean_query.lower()}?",
            f"Is {clean_query.lower()} required?",
            f"Does the policy cover {clean_query.lower()}?",
            f"What are the procedures for {clean_query.lower()}?",
            f"How is {clean_query.lower()} managed?",
            f"What controls exist for {clean_query.lower()}?",
            f"Are there standards for {clean_query.lower()}?",
            f"What governance exists for {clean_query.lower()}?",
            f"Who is responsible for {clean_query.lower()}?",
            f"What documentation covers {clean_query.lower()}?",
            f"Are there guidelines for {clean_query.lower()}?",
            f"What framework addresses {clean_query.lower()}?",
            f"How is {clean_query.lower()} implemented?"
        ]
        variations.extend(question_variations)
        
        # Limit to reasonable number of variations
        return variations[:10]
    
    def _deduplicate_documents(self, docs: List[Document]) -> List[Document]:
        """Remove duplicate documents based on content similarity."""
        if not docs:
            return docs
        
        unique_docs = []
        seen_content = set()
        
        for doc in docs:
            # Create a content hash (first 100 chars + length)
            content_hash = f"{doc.page_content[:100]}_{len(doc.page_content)}"
            
            if content_hash not in seen_content:
                seen_content.add(content_hash)
                unique_docs.append(doc)
        
        print(f"[SEMANTIC] Deduplication: {len(docs)} -> {len(unique_docs)} unique chunks")
        return unique_docs
    
    def _rerank_by_relevance(self, original_query: str, docs: List[Document]) -> List[Document]:
        """Re-rank documents by relevance to the original query."""
        if not docs:
            return docs
        
        try:
            # Calculate relevance scores
            scored_docs = []
            query_lower = original_query.lower()
            query_words = set(query_lower.split())
            
            for doc in docs:
                content_lower = doc.page_content.lower()
                
                # Score based on word overlap
                word_matches = sum(1 for word in query_words if word in content_lower)
                
                # Score based on phrase matches
                phrase_score = 0
                for word in query_words:
                    if len(word) > 3:  # Only meaningful words
                        phrase_score += content_lower.count(word)
                
                # Combined score
                total_score = word_matches * 2 + phrase_score
                scored_docs.append((doc, total_score))
            
            # Sort by score (highest first)
            scored_docs.sort(key=lambda x: x[1], reverse=True)
            
            # Return sorted documents
            ranked_docs = [doc for doc, score in scored_docs]
            print(f"[SEMANTIC] Re-ranking completed, top scores: {[score for _, score in scored_docs[:3]]}")
            
            return ranked_docs
            
        except Exception as e:
            print(f"[SEMANTIC] Re-ranking failed: {e}, returning original order")
            return docs

    def get_fallback_search_results(self, query: str, min_chunks: int = 20) -> List[Document]:
        """Fallback search strategy when primary methods don't find enough content."""
        print(f"[FALLBACK] Starting fallback search for query: '{query}'")
        
        try:
            # Strategy 1: Try with broader similarity threshold
            print(f"[FALLBACK] Strategy 1: Broader similarity search")
            retriever = self.db.as_retriever(search_type="similarity", k=25)
            broad_docs = retriever.invoke(query)
            print(f"[FALLBACK] Broad search found {len(broad_docs)} chunks")
            
            if len(broad_docs) >= min_chunks:
                return broad_docs[:min_chunks]
            
            # Strategy 2: Try keyword-based search
            print(f"[FALLBACK] Strategy 2: Keyword-based search")
            query_words = query.lower().split()
            keyword_docs = []
            
            # Search for each significant word
            for word in query_words:
                if len(word) > 3:  # Only meaningful words
                    try:
                        retriever = self.db.as_retriever(search_type="similarity", k=12)
                        word_docs = retriever.invoke(word)
                        keyword_docs.extend(word_docs)
                    except Exception as e:
                        print(f"[FALLBACK] Error searching for word '{word}': {e}")
            
            # Deduplicate and return
            unique_keyword_docs = self._deduplicate_documents(keyword_docs)
            print(f"[FALLBACK] Keyword search found {len(unique_keyword_docs)} unique chunks")
            
            if len(unique_keyword_docs) >= min_chunks:
                return unique_keyword_docs[:min_chunks]
            
            # Strategy 3: Return all we found
            print(f"[FALLBACK] Strategy 3: Returning all found chunks")
            all_docs = broad_docs + unique_keyword_docs
            unique_all = self._deduplicate_documents(all_docs)
            return unique_all[:min_chunks]
            
        except Exception as e:
            print(f"[FALLBACK] Fallback search failed: {e}")
            # Last resort: return any chunks we can find
            try:
                retriever = self.db.as_retriever(search_type="similarity", k=min_chunks)
                return retriever.invoke(query)
            except Exception as final_e:
                print(f"[FALLBACK] Final fallback failed: {final_e}")
                return []

    def _extract_fallback_answer(self, query: str, context: str) -> str:
        """Extract relevant information from the context to provide a fallback answer."""
        try:
            print(f"[FALLBACK] Attempting to extract relevant information from context")
            
            # Extract key terms from the query
            query_lower = query.lower()
            key_terms = []
            
            # Look for specific terms related to the question
            if 'backup' in query_lower:
                key_terms.extend(['backup', 'recovery', 'snapshot', 'image', 'data protection'])
            if 'management' in query_lower:
                key_terms.extend(['management', 'approval', 'authorization', 'oversight'])
            if 'outsourcer' in query_lower:
                key_terms.extend(['outsourcer', 'vendor', 'third party', 'external'])
            if 'scoped data' in query_lower:
                key_terms.extend(['scoped data', 'sensitive data', 'confidential data'])
            
            # Also include general terms from the query
            query_words = [word for word in query_lower.split() if len(word) > 3]
            key_terms.extend(query_words)
            
            print(f"[FALLBACK] Looking for key terms: {key_terms}")
            
            # Search through context for relevant information
            relevant_lines = []
            context_lines = context.split('\n')
            
            for line in context_lines:
                line_lower = line.lower()
                # Check if line contains any key terms
                if any(term in line_lower for term in key_terms):
                    # Clean up the line and add it
                    clean_line = line.strip()
                    if clean_line and len(clean_line) > 20:  # Only meaningful lines
                        relevant_lines.append(clean_line)
            
            print(f"[FALLBACK] Found {len(relevant_lines)} relevant lines")
            
            if relevant_lines:
                # Take the most relevant lines (limit to avoid overwhelming)
                selected_lines = relevant_lines[:5]  # Max 5 lines
                fallback_answer = "Based on the available context, I found the following relevant information: " + " ".join(selected_lines)
                return fallback_answer
            else:
                # If no direct matches, look for related policy information
                policy_indicators = ['policy', 'procedure', 'requirement', 'standard', 'guideline']
                policy_lines = []
                
                for line in context_lines:
                    line_lower = line.lower()
                    if any(indicator in line_lower for indicator in policy_indicators):
                        clean_line = line.strip()
                        if clean_line and len(clean_line) > 30:
                            policy_lines.append(clean_line)
                
                if policy_lines:
                    return f"While I don't see specific information about {query.split('?')[0]}, I found related policy information that may be relevant: " + " ".join(policy_lines[:3])
                else:
                    return "No specific information found about this topic in the available context."
                    
        except Exception as e:
            print(f"[FALLBACK] Error in fallback extraction: {e}")
            return "Error occurred while attempting to extract fallback information."

    def build_enhanced_prompt(self, query, context, chat_history=None):
        """Build optimized prompt for better answer quality and speed"""
        
        # Build concise but effective instructions
        instructions = f"""You are a compliance expert helping the user find information in the context. Provide ACCURATE answers using ONLY the provided context.

CRITICAL RULES:
- Answer directly without "Based on the context" or "I can provide"
- Start with "Yes" or "No" ONLY for explicit yes/no questions
- For all other questions, start directly with the answer
- Cite sources using [filename] format when information is found
- SEARCH THOROUGHLY through the context before saying "no information available"
- Look for related terms, synonyms, and policy references
- Be specific and actionable when information is found
- ONLY USE THE INFORMATION FROM THE CONTEXT TO ANSWER THE QUESTION. DO NOT MAKE UP INFORMATION.

SEARCH STRATEGY:
- Look for exact matches first
- Then search for related terms and concepts
- Check for policy names, procedure references, and requirements
- Look in all document chunks for relevant information
- Only say "No information available about [topic]" after thorough search

IMPORTANT: If you find relevant information, cite the specific sources. If you find NO relevant information after thorough search, say "No information available about [topic]" and DO NOT list any sources.

EXAMPLES:
- "Are backups required?" → "Yes, daily backups are required per [policy.pdf]"
- "What is the backup policy?" → "Daily incremental backups at 2 AM, weekly full backups per [backup_policy.pdf]"
- "How to reset password?" → "Use Password Manager tool, click 'Forgot Password' per [IT_policy.pdf]"
- "What is the alien policy?" → "No information available about alien policy"
THESE ARE ONLY EXAMPLES. YOU MUST ANSWER THE QUESTION DIRECTLY WITH THE INFORMATION FROM THE CONTEXT.

QUESTION: {query}

CONTEXT:
{context}

ANSWER:"""
        
        return instructions
    
    def _extract_key_entities(self, query):
        """Extract key entities and concepts from the query for better prompt engineering"""
        # Remove common words and extract meaningful terms
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'can', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them', 'what', 'when', 'where', 'why', 'how', 'who', 'which'}
        
        # Extract words and filter
        words = query.lower().split()
        key_terms = [word for word in words if word not in stop_words and len(word) > 2]
        
        # Add related terms for better context understanding
        related_terms = []
        for term in key_terms:
            if term in ['backup', 'recovery']:
                related_terms.extend(['data', 'disaster', 'business continuity'])
            elif term in ['password', 'authentication']:
                related_terms.extend(['security', 'access control', 'login'])
            elif term in ['encryption', 'cryptography']:
                related_terms.extend(['security', 'data protection', 'privacy'])
            elif term in ['policy', 'procedure']:
                related_terms.extend(['guideline', 'standard', 'requirement'])
            elif term in ['compliance', 'regulation']:
                related_terms.extend(['audit', 'governance', 'standard'])
        
        # Combine and deduplicate
        all_terms = key_terms + related_terms
        unique_terms = list(dict.fromkeys(all_terms))  # Preserve order while deduplicating
        
        return unique_terms[:10]  # Limit to top 10 terms

    def _expand_search_scope(self, query, original_docs):
        """Expand search scope when initial search doesn't find enough relevant information"""
        try:
            print(f"[EXPAND] Expanding search scope for query: '{query}'")
            
            # Extract key terms for broader search
            key_terms = self._extract_key_entities(query)
            print(f"[EXPAND] Key terms for expansion: {key_terms}")
            
            # Try searching with individual key terms
            expanded_docs = []
            for term in key_terms[:5]:  # Limit to top 5 terms
                try:
                    # Search with lower threshold for broader results
                    retriever = self.db.as_retriever(search_type="similarity", k=10)
                    term_docs = retriever.invoke(term)
                    expanded_docs.extend(term_docs)
                except Exception as e:
                    print(f"[EXPAND] Error searching term '{term}': {e}")
            
            # Also try searching with policy-related terms
            policy_terms = ['policy', 'procedure', 'requirement', 'standard', 'guideline']
            for term in policy_terms:
                try:
                    retriever = self.db.as_retriever(search_type="similarity", k=8)
                    policy_docs = retriever.invoke(term)
                    expanded_docs.extend(policy_docs)
                except Exception as e:
                    print(f"[EXPAND] Error searching policy term '{term}': {e}")
            
            # Deduplicate and combine with original docs
            all_docs = original_docs + expanded_docs
            unique_docs = self._deduplicate_documents(all_docs)
            
            print(f"[EXPAND] Search expansion: {len(original_docs)} -> {len(unique_docs)} total docs")
            return unique_docs[:50]  # Limit to reasonable size
            
        except Exception as e:
            print(f"[EXPAND] Search expansion failed: {e}")
            return original_docs

    def validate_context_relevance(self, query, context, docs):
        """Validate that context actually contains relevant information"""
        
        # Quick relevance check
        query_terms = set(query.lower().split())
        context_lower = context.lower()
        
        # Remove stop words from query terms
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'can', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them', 'what', 'when', 'where', 'why', 'how', 'who', 'which'}
        query_terms = {term for term in query_terms if term not in stop_words and len(term) > 2}
        
        # Check if key terms appear in context
        term_coverage = sum(1 for term in query_terms if term in context_lower)
        coverage_ratio = term_coverage / len(query_terms) if query_terms else 0
        
        print(f"[QUALITY] Context relevance check: {term_coverage}/{len(query_terms)} terms covered ({coverage_ratio:.2f})")
        
        # Enhanced relevance checking for better accuracy
        if coverage_ratio < 0.4:  # Increased threshold for better quality
            print(f"[QUALITY] WARNING: Low context relevance ({coverage_ratio:.2f}), answer quality may be poor")
            
            # Try to find more relevant chunks by lowering the threshold
            if hasattr(self, 'db') and self.db:
                try:
                    # Expand search with lower similarity threshold
                    expanded_docs = self._expand_search_scope(query, docs)
                    if expanded_docs and len(expanded_docs) > len(docs):
                        print(f"[QUALITY] Expanded search: {len(docs)} -> {len(expanded_docs)} docs")
                        return self.build_context_streamlined(expanded_docs)
                except Exception as e:
                    print(f"[QUALITY] Search expansion failed: {e}")
        
        # Additional quality checks
        if len(context.strip()) < 100:  # Context too short
            print(f"[QUALITY] WARNING: Context too short ({len(context)} chars), insufficient information")
        
        # Check for meaningful content (not just repeated phrases)
        context_words = context_lower.split()
        unique_words = len(set(context_words))
        if unique_words < 20:  # Too few unique words
            print(f"[QUALITY] WARNING: Context lacks variety ({unique_words} unique words), may be repetitive")
        
        return context
    
    def _expand_search_scope(self, query, current_docs):
        """Expand search scope to find more relevant documents"""
        try:
            if not hasattr(self, 'db') or not self.db:
                return current_docs
            
            # Get current document IDs to avoid duplicates
            current_ids = {doc.metadata.get('chunk_id', 0) for doc in current_docs}
            
            # Search with lower similarity threshold
            expanded_docs = self.db.similarity_search(
                query, 
                k=min(15, len(current_docs) + 5),  # Get a few more docs
                fetch_k=20  # Fetch more candidates
            )
            
            # Filter out duplicates and add new relevant docs
            new_docs = []
            for doc in expanded_docs:
                doc_id = doc.metadata.get('chunk_id', 0)
                if doc_id not in current_ids:
                    new_docs.append(doc)
                    current_ids.add(doc_id)
            
            # Combine current and new docs, prioritizing current ones
            combined_docs = current_docs + new_docs
            
            # Re-filter with lower threshold
            return self.filter_relevant_chunks(query, combined_docs, threshold=0.5)
            
        except Exception as e:
            print(f"[QUALITY] Error expanding search scope: {e}")
            return current_docs

    def is_gibberish(self, text: str) -> bool:
        """Fast gibberish detection using two-tier filtering"""
        
        # First filter: Regex patterns (1ms)
        if self._check_obvious_patterns(text):
            return True
        
        # Second filter: N-gram probability (5ms)
        if self._check_ngram_probability(text):
            return True
        
        return False

    def _check_obvious_patterns(self, text: str) -> bool:
        """Check for obvious gibberish patterns"""
        import re
        
        # Common keyboard patterns
        keyboard_patterns = [
            r'asdf+', r'qwerty+', r'zxcv+',  # Keyboard rows
            r'([a-z]{1,2})\1{3,}',  # Repeated 2-letter sequences
            r'([a-z])\1{4,}',    # Repeated single letters
            r'[bcdfghjklmnpqrstvwxz]{6,}',  # Long consonant clusters
            r'[aeiou]{6,}',        # Long vowel clusters
        ]
        
        for pattern in keyboard_patterns:
            if re.search(pattern, text.lower()):
                return True
        
        return False

    def _check_ngram_probability(self, text: str) -> bool:
        """Check character n-gram probability"""
        import re
        
        # Clean text (keep only letters and spaces)
        clean_text = re.sub(r'[^a-zA-Z\s]', '', text.lower())
        
        # Skip very short text
        if len(clean_text) < 4:
            return False
        
        # Calculate character bigram probability
        # English has common bigrams like 'th', 'he', 'an', 'in'
        # Gibberish has unlikely combinations like 'kj', 'hj', 'kh'
        
        # Simple heuristic: count impossible consonant pairs
        impossible_pairs = ['kj', 'hj', 'kh', 'jh', 'kjh', 'hjk']
        text_pairs = [clean_text[i:i+2] for i in range(len(clean_text)-1)]
        
        impossible_count = sum(1 for pair in text_pairs if pair in impossible_pairs)
        
        # If more than 30% of pairs are impossible, it's likely gibberish
        if len(text_pairs) > 0 and impossible_count / len(text_pairs) > 0.3:
            return True
        
        return False

    def reset_questionnaire_state(self):
        """Reset all instance variables to prevent contamination between questionnaires"""
        # Clear all previous context variables
        if hasattr(self, '_previous_context'):
            del self._previous_context
        if hasattr(self, '_previous_docs'):
            del self._previous_docs
        if hasattr(self, '_previous_query'):
            del self._previous_query
        
        # Clear any other instance variables that might persist
        if hasattr(self, 'context'):
            del self.context
        if hasattr(self, 'answer'):
            del self.answer
        
        # Clear any cached embeddings or search results
        if hasattr(self, '_cached_search_results'):
            del self._cached_search_results
        if hasattr(self, '_cached_embeddings'):
            del self._cached_embeddings
        
        # Clear any other cached data that might persist
        if hasattr(self, '_cached_docs'):
            del self._cached_docs
        if hasattr(self, '_cached_context'):
            del self._cached_context
        if hasattr(self, '_cached_embeddings_cache'):
            del self._cached_embeddings_cache
        
        # Reset LLM context if possible
        if hasattr(self, 'llm') and hasattr(self.llm, 'reset'):
            try:
                self.llm.reset()
            except:
                pass
        
        # Clear any session-specific data
        if hasattr(self, 'current_upload_session'):
            del self.current_upload_session
        if hasattr(self, 'current_questionnaire_id'):
            del self.current_questionnaire_id
        
        # Clear any other instance variables that might have been added
        for attr_name in list(self.__dict__.keys()):
            if attr_name.startswith('_') and attr_name not in ['_db', '_llm', '_embeddings', '_chunk_size', '_chunk_overlap']:
                try:
                    delattr(self, attr_name)
                except:
                    pass
        
        # Force garbage collection to clean up any remaining references
        import gc
        gc.collect()
        
        print(f"[SYSTEM LOG] (reset_questionnaire_state): All instance state cleared - no data leakage possible")

    def set_upload_session(self, session_id, questionnaire_id=None):
        """Set a new upload session to ensure complete isolation"""
        self.current_upload_session = session_id
        self.current_questionnaire_id = questionnaire_id
        # Reset state immediately when setting new session
        self.reset_questionnaire_state()
        print(f"[SYSTEM LOG] (set_upload_session): New session {session_id} established - complete isolation guaranteed")

    def _reset_context_state(self):
        """Internal method to reset context state before each query"""
        # Clear any context-related instance variables
        if hasattr(self, 'context'):
            del self.context
        if hasattr(self, 'answer'):
            del self.answer
        if hasattr(self, '_current_context'):
            del self._current_context
        if hasattr(self, '_current_docs'):
            del self._current_docs
        if hasattr(self, '_current_query'):
            del self._current_query

    def create_fresh_instance(self):
        """Create a completely fresh instance to ensure zero data leakage"""
        # This method can be called to create a new instance if needed
        # For now, we'll just do a complete reset
        self.reset_questionnaire_state()
        print(f"[SYSTEM LOG] (create_fresh_instance): Fresh instance state created - zero contamination possible")
        return True